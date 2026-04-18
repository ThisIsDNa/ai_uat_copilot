"""
DOCX export: business-ready UAT Review document (title page, scope & context, test results).

Session-state key patterns must stay aligned with `src/tc_session_keys.py` and
`src/ui_review_synthesis.py`.
"""

from __future__ import annotations

import re
from datetime import date
from io import BytesIO
from pathlib import Path
from typing import Any, Mapping

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

from src.coverage_gaps import generate_coverage_gaps
from src.intake_parser import load_scenario
from src.scenario_builder_core import format_tc_ac_link_lines, format_tc_ac_link_lines_for_export
from src.scenario_registry import build_scenario_catalog
from src.scenario_media import (
    expected_step_screenshot_paths,
    resolve_media_path,
    step_texts,
    resolved_test_case_title,
)
from src.traceability import generate_traceability_matrix
from src.tc_session_keys import (
    tc_flagged_review_key,
    tc_review_notes_key,
    tc_review_state_key,
    tc_review_upload_key,
    tc_row_session_suffixes,
)


def _test_results_status_key(scenario_key: str) -> str:
    return f"test_results_status_{scenario_key}"


def _test_results_status_reg_key(registry_id: str) -> str:
    return f"test_results_status_reg_{registry_id}"


def _session_test_results_status_label(
    session: Mapping[str, Any],
    *,
    scenario_key: str,
    registry_id: str | None,
) -> str:
    if registry_id:
        v = session.get(_test_results_status_reg_key(registry_id))
        if v is not None and _norm_str(v):
            return _norm_str(v)
    return _norm_str(session.get(_test_results_status_key(scenario_key))) or "In Progress"


def _norm_str(val: object) -> str:
    if val is None:
        return ""
    return str(val).strip()


def _changed_area_bullet_lines(raw: object) -> list[str]:
    """Normalize ``changed_areas`` to non-empty bullet strings (list of str or dict rows)."""
    if raw is None:
        return []
    if isinstance(raw, str):
        s = raw.strip()
        return [s] if s else []
    if not isinstance(raw, list):
        return []
    out: list[str] = []
    for row in raw:
        if isinstance(row, dict):
            parts = [f"{k}: {_norm_str(v)}" for k, v in row.items() if _norm_str(v)]
            line = "; ".join(parts) if parts else ""
            if line.strip():
                out.append(line.strip())
        else:
            line = _norm_str(row)
            if line:
                out.append(line)
    return out


_EXEC_DRAFT_NON_FORM_TITLE_HINTS: tuple[str, ...] = (
    "reply",
    "draft",
    "generate",
    "unauthorized",
    "permission denied",
    "service failure",
    "ai ",
    " ai",
    "token",
    "session",
    "blocked",
    "precondition",
    "persist",
    "refresh",
    "failure",
    "rejection",
    "validation error",
)
_EXEC_DRAFT_GENERIC_STEP_RES: tuple[re.Pattern[str], ...] = (
    re.compile(r"^update\s+(all\s+)?required\s+field", re.I),
    re.compile(r"^enter\s+.*required\s+field", re.I),
    re.compile(r"^populate\s+(all\s+)?required\s+field", re.I),
    re.compile(r"^ensure\s+all\s+required\s+field", re.I),
    re.compile(r"^complete\s+all\s+required\s+field", re.I),
    re.compile(r"^fill\s+in\s+all\s+required\s+field", re.I),
)


def _execution_draft_title_suggests_non_form_workflow(title: str) -> bool:
    tl = (title or "").lower()
    return any(h in tl for h in _EXEC_DRAFT_NON_FORM_TITLE_HINTS)


def _strip_leaky_generic_steps_for_execution_draft(title: str, steps: list[str]) -> list[str]:
    """
    Drop generic required-field boilerplate from exports when the test case title clearly targets
    another intent (reply/draft/AI failure paths, etc.). Never drops every step.
    """
    if not steps or not _execution_draft_title_suggests_non_form_workflow(title):
        return steps
    kept: list[str] = []
    for s in steps:
        st = _norm_str(s)
        if not st:
            continue
        if any(rx.match(st) for rx in _EXEC_DRAFT_GENERIC_STEP_RES):
            continue
        kept.append(s)
    return kept if kept else steps


def _known_dependency_bullet_lines(raw: object) -> list[str]:
    """Normalize ``known_dependencies`` to non-empty bullet strings."""
    if raw is None:
        return []
    if isinstance(raw, str):
        s = raw.strip()
        return [s] if s else []
    if not isinstance(raw, list):
        return []
    out: list[str] = []
    for d in raw:
        line = _norm_str(d)
        if line:
            out.append(line)
    return out


def _evidence_basename(relative_path: str) -> str:
    """File name only for stakeholder-facing labels (no internal project paths)."""
    r = _norm_str(relative_path).replace("\\", "/")
    if not r:
        return ""
    return Path(r).name or r


def _figure_step_evidence_caption(figure_n: int, step_one_based: int, relative_path: str) -> str:
    """e.g. Figure 1 - Step 1 - tc_01_step_01.png"""
    fname = _evidence_basename(relative_path) or "screenshot.png"
    return f"Figure {figure_n} - Step {step_one_based} - {fname}"


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_para(
    doc: Document, text: str, *, style: str | None = None, space_after: Pt | None = None
) -> Paragraph:
    p = doc.add_paragraph(text, style=style) if style else doc.add_paragraph(text)
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    return p


def _add_spacer(doc: Document, pt: int = 12) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)


def _add_section_divider(doc: Document) -> None:
    """Thin horizontal rule between major sections."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(18)
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFBFBF")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _title_page(
    doc: Document,
    data: dict,
    session: Mapping[str, Any],
    *,
    scenario_key: str,
    registry_id: str | None,
) -> None:
    story = (
        _norm_str(data.get("scenario_title"))
        or _norm_str(data.get("story_title"))
        or _norm_str(data.get("scenario_id"))
        or scenario_key
    )
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(8)
    r = p_title.add_run(story)
    r.bold = True
    r.font.size = Pt(26)

    bg = _norm_str(data.get("business_goal"))
    if bg:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_after = Pt(24)
        r2 = p_sub.add_run(bg)
        r2.font.size = Pt(13)

    _add_spacer(doc, 18)

    sid = _norm_str(data.get("scenario_id")) or scenario_key
    wf = _norm_str(data.get("workflow_name"))
    status = _session_test_results_status_label(
        session, scenario_key=scenario_key, registry_id=registry_id
    )
    export_d = date.today().isoformat()

    _add_para(doc, f"Scenario ID: {sid}", space_after=Pt(6))
    if wf:
        _add_para(doc, f"Workflow Name: {wf}", space_after=Pt(6))
    _add_para(doc, f"Review Status: {status}", space_after=Pt(6))
    _add_para(doc, f"Export Date: {export_d}", space_after=Pt(6))

    _add_spacer(doc, 24)


def _scope_context_section(doc: Document, data: dict) -> None:
    _add_heading(doc, "Scope & Context", level=1)
    _add_spacer(doc, 6)

    changed_lines = _changed_area_bullet_lines(data.get("changed_areas"))
    if changed_lines:
        _add_heading(doc, "Changed Areas", level=2)
        for line in changed_lines:
            doc.add_paragraph(line, style="List Bullet")
        _add_spacer(doc, 12)

    dep_lines = _known_dependency_bullet_lines(data.get("known_dependencies"))
    if dep_lines:
        _add_heading(doc, "Known Dependencies", level=2)
        for line in dep_lines:
            doc.add_paragraph(line, style="List Bullet")
        _add_spacer(doc, 12)

    _add_heading(doc, "Acceptance Criteria", level=2)
    acs = [x for x in (data.get("acceptance_criteria") or []) if isinstance(x, dict)]
    if acs:
        for ac in acs:
            ac_id = _norm_str(ac.get("id")) or "—"
            text = _norm_str(ac.get("text")) or "—"
            doc.add_paragraph(f"{ac_id}: {text}", style="List Bullet")
    else:
        _add_para(doc, "No acceptance criteria in this scenario.")
    _add_spacer(doc, 12)


def _gap_description_bullets(desc: str) -> list[str]:
    t = desc.strip()
    if not t:
        return []
    if "|" not in t:
        return [t]
    return [p.strip() for p in t.split("|") if p.strip()]


def _coverage_gaps_subsection(doc: Document, gap_rows: list[dict]) -> None:
    _add_heading(doc, "Coverage Gaps", level=2)
    _add_spacer(doc, 6)

    if not gap_rows:
        _add_para(doc, "No coverage gaps were flagged for this scenario.")
        _add_spacer(doc, 12)
        return

    by_ac: dict[str, list[dict]] = {}
    order: list[str] = []
    for row in gap_rows:
        if not isinstance(row, dict):
            continue
        ac = _norm_str(row.get("acceptance_criteria_id")) or "—"
        if ac not in by_ac:
            by_ac[ac] = []
            order.append(ac)
        by_ac[ac].append(row)

    for ac in order:
        for i, row in enumerate(by_ac[ac]):
            if i > 0:
                _add_spacer(doc, 8)
            gt = _norm_str(row.get("gap_type")) or "—"
            desc_raw = _norm_str(row.get("description")) or "—"
            act = _norm_str(row.get("suggested_action"))

            _add_heading(doc, f"{ac} — {gt}", level=3)
            for frag in _gap_description_bullets(desc_raw):
                doc.add_paragraph(frag, style="List Bullet")
            if act:
                _add_para(doc, f"Suggested action: {act}", space_after=Pt(4))

    _add_spacer(doc, 12)


def _try_add_picture(doc: Document, raw: bytes, caption: str) -> bool:
    if not raw:
        return False
    bio = BytesIO(raw)
    try:
        doc.add_picture(bio, width=Inches(5.2))
        if caption:
            cap = doc.add_paragraph(caption)
            cap.paragraph_format.space_after = Pt(10)
        return True
    except Exception:
        pass
    try:
        from PIL import Image

        bio.seek(0)
        img = Image.open(bio)
        rgb = img.convert("RGB")
        out = BytesIO()
        rgb.save(out, format="PNG")
        out.seek(0)
        doc.add_picture(out, width=Inches(5.2))
        if caption:
            cap = doc.add_paragraph(caption)
            cap.paragraph_format.space_after = Pt(10)
        return True
    except Exception:
        _add_para(doc, f"{caption} (image could not be embedded.)")
        return False


def _test_results_subsection(
    doc: Document,
    *,
    test_cases: list,
    scenario_key: str,
    session: Mapping[str, Any],
    data: dict,
) -> None:
    _add_heading(doc, "Test Results", level=2)
    _add_spacer(doc, 6)

    suffixes = tc_row_session_suffixes(test_cases)
    included: list[tuple[str, dict]] = []
    for i, tc in enumerate(test_cases):
        if not isinstance(tc, dict):
            continue
        row_suf = suffixes[i]
        reviewed = bool(session.get(tc_review_state_key(scenario_key, row_suf), False))
        flagged = bool(session.get(tc_flagged_review_key(scenario_key, row_suf), False))
        if not reviewed and not flagged:
            continue
        included.append((row_suf, tc))

    if not included:
        _add_para(
            doc,
            "No test cases are included in this export. Only cases marked "
            "**Flagged for Review** or **Reviewed & Approved** in the app are listed here.",
        )
        return

    figure_n = 0
    for row_suf, tc in included:
        tid_str = _norm_str(tc.get("id")) or "unknown"
        tc_text = _norm_str(resolved_test_case_title(tc)) or "—"
        reviewed = bool(session.get(tc_review_state_key(scenario_key, row_suf), False))
        flagged = bool(session.get(tc_flagged_review_key(scenario_key, row_suf), False))

        if flagged and reviewed:
            status_line = "Flagged for Review; Reviewed & Approved"
        elif flagged:
            status_line = "Flagged for Review"
        else:
            status_line = "Reviewed & Approved"

        _add_heading(doc, f"{tid_str}: {tc_text}", level=3)
        ac_lines = format_tc_ac_link_lines(data.get("acceptance_criteria"), tid_str)
        if ac_lines:
            _add_para(doc, "Linked acceptance criteria:", space_after=Pt(4))
            for al in ac_lines:
                plain = al.replace("**", "")
                doc.add_paragraph(plain, style="List Bullet")
            _add_spacer(doc, 4)
        _add_para(doc, f"Status: {status_line}", space_after=Pt(8))

        author_notes = _norm_str(tc.get("notes")) if isinstance(tc, dict) else ""
        if author_notes:
            _add_para(doc, "Tester notes (from scenario author):", space_after=Pt(4))
            for ln in author_notes.splitlines():
                ln = ln.strip()
                if ln:
                    doc.add_paragraph(ln, style="List Bullet")
            _add_spacer(doc, 6)

        steps = step_texts(tc)
        paths = expected_step_screenshot_paths(tc)
        nshow = max(len(steps), len(paths))
        if nshow > 0:
            _add_para(doc, "Steps and expected step evidence", space_after=Pt(4))
            for si in range(nshow):
                stxt = _norm_str(steps[si]) if si < len(steps) else ""
                rel = _norm_str(paths[si]) if si < len(paths) else ""
                line = f"Step {si + 1}: {stxt}" if stxt else f"Step {si + 1}: (no step text)"
                _add_para(doc, line, space_after=Pt(2))
                if rel:
                    figure_n += 1
                    cap = _figure_step_evidence_caption(figure_n, si + 1, rel)
                    rp = resolve_media_path(rel)
                    if rp is not None and rp.is_file():
                        try:
                            raw_img = rp.read_bytes()
                            _try_add_picture(doc, raw_img, cap)
                        except OSError:
                            _add_para(
                                doc,
                                f"Step {si + 1}: expected evidence could not be read ({cap}).",
                                space_after=Pt(2),
                            )
                    else:
                        _add_para(
                            doc,
                            f"Step {si + 1}: expected evidence file not available ({cap}).",
                            space_after=Pt(2),
                        )
                else:
                    _add_para(
                        doc,
                        f"Step {si + 1}: no expected evidence reference for this step.",
                        space_after=Pt(2),
                    )
            _add_spacer(doc, 8)

        notes = session.get(tc_review_notes_key(scenario_key, row_suf))
        notes_s = _norm_str(notes) if notes is not None else ""
        if notes_s:
            _add_para(doc, "Reviewer Notes:", space_after=Pt(4))
            for line in notes_s.splitlines():
                line = line.strip()
                if line:
                    doc.add_paragraph(line, style="List Bullet")

        up = session.get(tc_review_upload_key(scenario_key, row_suf))
        raw: bytes | None = None
        fname = ""
        if up is not None and hasattr(up, "getvalue"):
            try:
                raw = up.getvalue()
            except Exception:
                raw = None
            fname = _norm_str(getattr(up, "name", ""))

        if raw:
            figure_n += 1
            _add_para(doc, "Supporting Evidence:", space_after=Pt(4))
            upload_bn = _evidence_basename(fname) if fname else ""
            cap = f"Figure {figure_n} - Supporting evidence"
            if upload_bn:
                cap = f"{cap} - {upload_bn}"
            _try_add_picture(doc, raw, cap)

        _add_spacer(doc, 14)


def _test_results_main_section(
    doc: Document,
    *,
    test_cases: list,
    scenario_key: str,
    session: Mapping[str, Any],
    gap_rows: list[dict],
    data: dict,
) -> None:
    _add_heading(doc, "Test Results", level=1)
    _add_spacer(doc, 8)
    _coverage_gaps_subsection(doc, gap_rows)
    _add_spacer(doc, 6)
    _test_results_subsection(
        doc,
        test_cases=test_cases,
        scenario_key=scenario_key,
        session=session,
        data=data,
    )


def build_uat_review_export_docx(
    *,
    data: dict,
    test_cases: list,
    scenario_key: str,
    traceability_matrix: list[dict],
    session: Mapping[str, Any],
    tc_to_explicit_acs: dict[str, list[str]],
    gap_rows: list[dict] | None = None,
    reviewer_focus: dict[str, list[str]] | None = None,
    registry_id: str | None = None,
) -> bytes:
    """
    Build a presentation-ready UAT Review DOCX from scenario data and session values.
    Structure: Title Page → Scope & Context → Test Results (coverage gaps + per-TC outcomes).
    """
    _ = tc_to_explicit_acs, reviewer_focus  # API compatibility with callers

    # Saved scenarios: always read the latest JSON from disk so export matches Scenario Editor saves.
    rid = str(registry_id or "").strip()
    if rid:
        try:
            meta = build_scenario_catalog().get(rid)
            if isinstance(meta, dict):
                pth = _norm_str(meta.get("path"))
                if pth:
                    loaded = load_scenario(pth)
                    data = loaded
                    test_cases = list(loaded.get("test_cases") or [])
                    traceability_matrix = generate_traceability_matrix(loaded)
                    gap_rows = None
        except Exception:
            pass

    gaps = gap_rows
    if gaps is None:
        try:
            gaps = generate_coverage_gaps(data, traceability_rows=traceability_matrix)
        except Exception:
            gaps = []

    doc = Document()

    _title_page(doc, data, session, scenario_key=scenario_key, registry_id=registry_id)
    _add_section_divider(doc)
    _scope_context_section(doc, data)
    _add_section_divider(doc)
    _test_results_main_section(
        doc,
        test_cases=test_cases,
        scenario_key=scenario_key,
        session=session,
        gap_rows=gaps or [],
        data=data,
    )

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def safe_export_filename(data: dict, scenario_key: str) -> str:
    base = _norm_str(data.get("scenario_id")) or scenario_key or "scenario"
    safe = re.sub(r"[^\w.\-]+", "_", base, flags=re.UNICODE).strip("._") or "scenario"
    return f"uat_export_{safe}.docx"


def safe_execution_draft_filename(data: dict) -> str:
    """Download name for pre-review execution draft exports (distinct from final UAT export)."""
    base = _norm_str(data.get("scenario_id")) or "scenario"
    safe = re.sub(r"[^\w.\-]+", "_", base, flags=re.UNICODE).strip("._") or "scenario"
    return f"{safe}_execution_draft.docx"


def _execution_draft_title_block(doc: Document, data: dict) -> None:
    p_banner = doc.add_paragraph()
    p_banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rb = p_banner.add_run("Execution draft — draft test script (pre-review)")
    rb.bold = True
    rb.font.size = Pt(13)

    p_tag = doc.add_paragraph()
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = p_tag.add_run("Pre-review export · not an approved final UAT record")
    rt.italic = True
    rt.font.size = Pt(10)
    p_tag.paragraph_format.space_after = Pt(14)

    story = (
        _norm_str(data.get("scenario_title"))
        or _norm_str(data.get("story_title"))
        or _norm_str(data.get("scenario_id"))
        or "Untitled scenario"
    )
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(8)
    r = p_title.add_run(story)
    r.bold = True
    r.font.size = Pt(24)

    sid = _norm_str(data.get("scenario_id")) or "—"
    wf = _norm_str(data.get("workflow_name"))
    bg = _norm_str(data.get("business_goal"))
    export_d = date.today().isoformat()

    _add_para(doc, f"Scenario ID: {sid}", space_after=Pt(4))
    if wf:
        _add_para(doc, f"Workflow: {wf}", space_after=Pt(4))
    if bg:
        _add_para(doc, f"Business goal: {bg}", space_after=Pt(4))
    _add_para(doc, f"Export date: {export_d}", space_after=Pt(12))


def _execution_draft_scenario_context_section(doc: Document, data: dict) -> None:
    _add_heading(doc, "Scenario context", level=1)
    _add_spacer(doc, 4)
    ctx = _norm_str(data.get("scenario_context") or data.get("story_description") or "")
    if ctx:
        for para in ctx.splitlines():
            line = para.strip()
            if line:
                _add_para(doc, line, space_after=Pt(4))
    else:
        _add_para(doc, "(No scenario context text in this draft.)")
    _add_spacer(doc, 12)

    changed_lines = _changed_area_bullet_lines(data.get("changed_areas"))
    if changed_lines:
        _add_heading(doc, "Changed Areas", level=1)
        _add_spacer(doc, 4)
        for line in changed_lines:
            doc.add_paragraph(line, style="List Bullet")
        _add_spacer(doc, 12)

    dep_lines = _known_dependency_bullet_lines(data.get("known_dependencies"))
    if dep_lines:
        _add_heading(doc, "Known Dependencies", level=1)
        _add_spacer(doc, 4)
        for line in dep_lines:
            doc.add_paragraph(line, style="List Bullet")
        _add_spacer(doc, 12)


def _execution_draft_acceptance_and_scope(doc: Document, data: dict) -> None:
    _add_heading(doc, "Acceptance criteria", level=1)
    _add_spacer(doc, 4)
    acs = [x for x in (data.get("acceptance_criteria") or []) if isinstance(x, dict)]
    if acs:
        for ac in acs:
            ac_id = _norm_str(ac.get("id")) or "—"
            text = _norm_str(ac.get("text")) or "—"
            doc.add_paragraph(f"{ac_id}: {text}", style="List Bullet")
    else:
        _add_para(doc, "No acceptance criteria in this draft.")
    _add_spacer(doc, 12)


def _execution_draft_test_cases_section(doc: Document, data: dict) -> None:
    _add_heading(doc, "Test cases & steps", level=1)
    _add_spacer(doc, 6)
    test_cases = [x for x in (data.get("test_cases") or []) if isinstance(x, dict)]
    acs = data.get("acceptance_criteria") or []
    if not test_cases:
        _add_para(doc, "No test cases in this draft.")
        return

    for tc in test_cases:
        tid = _norm_str(tc.get("id")) or "—"
        title = _norm_str(resolved_test_case_title(tc)) or "—"
        _add_heading(doc, f"{tid}: {title}", level=2)

        pref_slot = tc.get("_export_primary_ac_slot")
        pref_i: int | None = None
        if pref_slot is not None and not isinstance(pref_slot, bool):
            try:
                pref_i = int(pref_slot)
            except (TypeError, ValueError):
                pref_i = None
        ac_lines = format_tc_ac_link_lines_for_export(
            acs if isinstance(acs, list) else [],
            tid,
            prefer_ac_slot_index=pref_i,
        )
        if ac_lines:
            _add_para(doc, "Linked acceptance criteria:", space_after=Pt(4))
            for al in ac_lines:
                plain = al.replace("**", "")
                doc.add_paragraph(plain, style="List Bullet")
            _add_spacer(doc, 4)

        steps = _strip_leaky_generic_steps_for_execution_draft(title, list(step_texts(tc)))
        _add_heading(doc, "Test steps", level=3)
        if steps:
            for si, st in enumerate(steps):
                stxt = _norm_str(st) if st else "(no step text)"
                doc.add_paragraph(f"{si + 1}. {stxt}")
        else:
            _add_para(doc, "(No steps drafted yet.)")

        author_notes = _norm_str(tc.get("notes"))
        if author_notes:
            _add_spacer(doc, 6)
            _add_para(doc, "Tester notes (from builder):", space_after=Pt(4))
            for ln in author_notes.splitlines():
                ln = ln.strip()
                if ln:
                    doc.add_paragraph(ln, style="List Bullet")

        _add_spacer(doc, 8)
        _add_heading(doc, "Execution", level=3)
        _add_para(doc, "After running the steps above:", space_after=Pt(6))
        _add_para(
            doc,
            "Test result (Pass / Fail / Blocked): _______________________________________________",
            space_after=Pt(4),
        )
        _add_para(
            doc,
            "Actual result / notes: _______________________________________________",
            space_after=Pt(4),
        )
        _add_para(
            doc,
            "Evidence / screenshot (path, file name, or link): _______________________________________________",
            space_after=Pt(10),
        )

        _add_spacer(doc, 16)


def build_execution_draft_export_docx(*, data: dict, test_cases: list | None = None) -> bytes:
    """
    Pre-review execution packet: current scenario content + fill-in placeholders for QA.

    Does not use review session flags, coverage gates, or registry reload — callers pass
    the authoritative dict (e.g. ``read_flat_builder_session``).
    """
    tcs = list(test_cases if test_cases is not None else (data.get("test_cases") or []))
    data = dict(data)
    data["test_cases"] = tcs

    doc = Document()
    _execution_draft_title_block(doc, data)
    _add_section_divider(doc)
    _execution_draft_scenario_context_section(doc, data)
    _add_section_divider(doc)
    _execution_draft_acceptance_and_scope(doc, data)
    _add_section_divider(doc)
    _execution_draft_test_cases_section(doc, data)

    notes = _norm_str(data.get("notes"))
    if notes:
        _add_section_divider(doc)
        _add_heading(doc, "Scenario notes (author)", level=1)
        _add_spacer(doc, 4)
        for ln in notes.splitlines():
            ln = ln.strip()
            if ln:
                _add_para(doc, ln, space_after=Pt(4))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
