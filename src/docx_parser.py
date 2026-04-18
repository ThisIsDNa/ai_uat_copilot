"""
Parse UAT-style Word (.docx) documents into the same scenario dict shape as JSON scenarios.

Assumptions (see project docs for schema):
- Labeled lines and Word heading styles introduce sections; loose headings are inferred
  (requirements, validation criteria, goal/purpose, scope, TC ids).
- Test cases: "TC-01", "TC 01", "Test case 1", "Test Case: ...", "Test: ..." (title
  after the colon), "Case 1:"; a heading line that is itself a TC id opens that case.
- Acceptance criteria: section headers, AC-/FR-/REQ-style lines and tables, GWT lines,
  and (before any TC exists) numbered or bulleted lines that read like criteria—not
  imperative UAT steps.
- Steps: "1." / "1)" / "a)" / "Step 1:", Word numbered list items in the tests section,
  or table rows (Step | Expected) in test areas.
- Lines like "Expected result: ..." merge into the previous step text (schema stays
  one string per step).
- Bullets (•, -, *) are stripped; acceptance vs test steps is decided by context.
- Images: extracted from drawing/VML; optional **Figure / Fig / [Screenshot N]** captions
  register labels; step text like **see Figure 1.1** maps images to steps. Unreferenced
  images still use order-based placement; extras go to workflow_process_screenshots.

Debug:
- Logging: set level DEBUG for ``src.docx_parser`` (section headers, table rows, AC/TC linking).
- Structure dump (CLI): ``python scripts/inspect_docx_structure.py <file.docx>`` — headings,
  paragraph styles, table row previews, and which lines match section/test-case patterns.
"""

from __future__ import annotations

import hashlib
import io
import logging
import re
from pathlib import Path
from typing import BinaryIO

logger = logging.getLogger(__name__)

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.scenario_media import (
    raw_step_screenshot_paths_in_json_order,
    resolved_step_screenshots,
)

_PROJECT_ROOT = Path(__file__).resolve().parent.parent
_IMPORT_ROOT = _PROJECT_ROOT / "data" / "docx_imports"

# Wordprocessing ML drawing inline (image in run)
_WP_DRAWING = (
    "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
)
_DML_MAIN = "http://schemas.openxmlformats.org/drawingml/2006/main"
_VML_NS = "urn:schemas-microsoft-com:vml"
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _slug_from_name(name: str, content_hash: str) -> str:
    stem = Path(name).stem if name else "upload"
    safe = re.sub(r"[^\w\-]+", "_", stem, flags=re.UNICODE).strip("_")[:40]
    return f"{safe}_{content_hash[:8]}" if safe else f"import_{content_hash[:8]}"


def _content_hash(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def iter_block_items(parent: DocumentObject):
    """Yield Paragraph and Table children in document order."""
    body = parent.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _ext_for_image_part(part) -> str:
    ct = (getattr(part, "content_type", None) or "").lower()
    if "png" in ct:
        return ".png"
    if "jpeg" in ct or "jpg" in ct:
        return ".jpg"
    if "gif" in ct:
        return ".gif"
    if "bmp" in ct:
        return ".bmp"
    return ".bin"


def _collect_r_embed_ids_from_paragraph(paragraph: Paragraph) -> list[str]:
    """All unique r:embed ids for images under this paragraph (any depth)."""
    root = paragraph._element
    seen: set[str] = set()
    ordered: list[str] = []
    for blip in root.findall(f".//{{{_DML_MAIN}}}blip"):
        rid = blip.get(qn("r:embed"))
        if rid and rid not in seen:
            seen.add(rid)
            ordered.append(rid)
    for im in root.findall(f".//{{{_VML_NS}}}imagedata"):
        rid = im.get(qn("r:id")) or im.get(qn("r:embed"))
        if rid and rid not in seen:
            seen.add(rid)
            ordered.append(rid)
    return ordered


def _save_paragraph_images(
    paragraph: Paragraph,
    out_dir: Path,
    image_index: list[int],
) -> list[str]:
    """
    Extract images from a paragraph (drawing + VML); preserve document order of embeds.
    Returns paths relative to project root (posix-style).
    """
    saved: list[str] = []
    part = paragraph.part
    for r_embed in _collect_r_embed_ids_from_paragraph(paragraph):
        try:
            image_part = part.related_parts[r_embed]
        except KeyError:
            continue
        blob = getattr(image_part, "blob", None)
        if not blob:
            continue
        ext = _ext_for_image_part(image_part)
        image_index[0] += 1
        fname = f"img_{image_index[0]:04d}{ext}"
        fpath = out_dir / fname
        fpath.write_bytes(blob)
        rel = fpath.relative_to(_PROJECT_ROOT).as_posix()
        saved.append(rel)
    return saved


def _is_heading_paragraph(paragraph: Paragraph) -> bool:
    try:
        n = (paragraph.style.name or "").strip().lower()
        if n.startswith("heading"):
            return True
        if n in ("title", "subtitle"):
            return True
    except (AttributeError, TypeError):
        pass
    return False


def _paragraph_has_numbering(paragraph: Paragraph) -> bool:
    p = paragraph._element
    ppr = p.find(f"{{{_W_NS}}}pPr")
    if ppr is None:
        return False
    return ppr.find(f"{{{_W_NS}}}numPr") is not None


def _normalize_line_for_parse(line: str) -> str:
    """Strip common bullet / dash prefixes (body text, not headings)."""
    s = line.strip()
    if not s:
        return s
    s = re.sub(
        r"^[\s\u00b7\u2022\u2023\u25aa\u25cf\u2043\u2219\-\*•·]+",
        "",
        s,
    )
    return s.strip()


def _is_tests_subsection_label(line: str) -> bool:
    t = _normalize_header(line)
    if len(t) > 60:
        return False
    return bool(
        re.match(
            r"^(test\s*)?steps?\s*:?$|^(procedure|actions?|execution)\s*:?$",
            t,
        )
        or re.match(r"^expected\s*(results?|outcomes?)?\s*:?$", t)
    )


def _merge_expected_result_into_last_step(state: "_DocxParseState", detail: str) -> bool:
    """Append expected-result text to the last step of the current TC. Returns True if applied."""
    detail = detail.strip()
    if not detail or not state.test_cases:
        return False
    steps = state.test_cases[-1].get("steps") or []
    if not steps:
        return False
    sep = " — Expected: "
    state.test_cases[-1]["steps"][-1] = steps[-1].rstrip() + sep + detail
    return True


def _normalize_header(line: str) -> str:
    return re.sub(r"\s+", " ", line.strip().lower())


def _strip_leading_outline(t: str) -> str:
    """Remove '1.', '1.1.', 'Section 2:' style prefixes from normalized headers."""
    t = t.strip()
    t = re.sub(r"^\d+(\.\d+)*[\.\)]\s*", "", t)
    t = re.sub(r"^(section|part)\s+\d+\s*[:.\-]?\s*", "", t)
    return t.strip()


def _match_section_header(text: str) -> str | None:
    t = _normalize_header(text)
    t = _strip_leading_outline(t)
    if len(t) > 80:
        return None
    # Test case id as its own heading line — require ':'/'-' title or line ends after id
    # so prose like "TC-01 validation must pass" stays out of this branch.
    if re.match(r"^tc[\s\-]+\d+\s*$", t) or re.match(
        r"^tc[\s\-]+\d+\s*[:\-–\.]\s*\S", t
    ):
        return "tests"
    if re.match(r"^(story\s*)?title\s*:", t) or t == "title":
        return "story_title"
    if re.match(r"^(story\s*)?(description|overview)\s*:", t) or t in (
        "description",
        "overview",
    ):
        return "story_description"
    if re.match(r"^user\s*story\s*:", t) or t == "user story":
        return "story_description"
    if re.match(r"^scenario\s*:", t) and "acceptance" not in t:
        return "story_description"
    if (
        re.match(r"^business\s*goals?\s*:", t)
        or re.match(r"^business\s*goals?$", t)
        or t.startswith("business goal")
        or t.startswith("business goals")
        or re.match(r"^business\s*objectives?\s*:?", t)
    ):
        return "business_goal"
    if re.match(r"^(goal|goals|purpose|objectives?)\s*:", t) or t in (
        "goal",
        "goals",
        "purpose",
        "objective",
        "objectives",
    ):
        return "business_goal"
    if "objective" in t and len(t) < 55 and "acceptance" not in t:
        return "business_goal"
    if re.match(r"^strategic\s*(goal|objective)s?\s*:?", t) and len(t) < 70:
        return "business_goal"
    if ("acceptance" in t and "criter" in t) or t in (
        "acceptance",
        "requirements",
        "functional requirements",
        "user requirements",
        "success criteria",
        "validation criteria",
        "product requirements",
        "business requirements",
    ):
        return "acceptance"
    if t in ("success criteria", "exit criteria", "definition of done", "sign off criteria"):
        return "acceptance"
    if "definition of done" in t and len(t) < 60:
        return "acceptance"
    if re.match(r"^user\s*acceptance\b", t) and len(t) < 70:
        return "acceptance"
    if re.match(r"^test\s*(cases?|scenarios?|scripts?|plans?|steps?)\s*:?$", t):
        return "tests"
    if t in (
        "uat",
        "uat tests",
        "uat test cases",
        "uat scenarios",
        "qa tests",
        "qa test cases",
        "testing",
        "quality assurance",
        "test execution",
        "test plan",
    ):
        return "tests"
    if "test" in t and "case" in t and len(t) < 55:
        return "tests"
    if re.match(r"^test\s*cases?\s*:", t) or t == "test cases":
        return "tests"
    # Narrative / template heading — must not route to changed_areas (was leaking "& Context" rows).
    if re.match(r"^scope\s*(?:&|and)\s*context\b", t) and len(t) < 70:
        return "workflow_ctx"
    if ("changed" in t and "area" in t) or t in (
        "scope of change",
        "scope",
        "in scope",
        "out of scope",
        "impacted areas",
        "impact analysis",
        "modules affected",
        "affected areas",
        "areas changed",
    ):
        return "changed_areas"
    if re.match(r"^(in[-\s]?scope|out[-\s]?of[-\s]?scope)\s*:?$", t):
        return "changed_areas"
    if t in (
        "changes",
        "change summary",
        "what changed",
        "components affected",
        "touchpoints",
        "systems affected",
        "impacted components",
        "impacted systems",
    ) or (len(t) < 55 and "impact" in t and ("area" in t or "component" in t)):
        return "changed_areas"
    if (len(t) < 40 and t.startswith("change") and "area" in t):
        return "changed_areas"
    if ("depend" in t and len(t) < 70) or "prerequisite" in t or "related system" in t:
        return "dependencies"
    if re.search(
        r"\b(integrations?|interfaces?|external\s+systems?|system\s+interfaces?|"
        r"dependent\s+systems?|upstream|downstream|data\s+feeds?)\b",
        t,
    ) and len(t) < 85:
        return "dependencies"
    if re.match(r"^workflow", t) or re.match(r"^context", t):
        return "workflow_ctx"
    if re.search(
        r"\b(background|clinical\s+workflow|patient\s+(flow|journey)|end[-\s]?to[-\s]?end|"
        r"process\s+(flow|overview)|application\s+flow|operational\s+context|"
        r"current\s+state|as[-\s]?is\s+process)\b",
        t,
    ) and len(t) < 90:
        return "workflow_ctx"
    if re.search(
        r"\b(modifications?|updates?\s+to|altered?\s+components?|release\s+impact|"
        r"affected\s+applications?|systems?\s+impacted)\b",
        t,
    ) and len(t) < 85:
        return "changed_areas"
    if "application" in t and "impact" in t and len(t) < 85:
        return "changed_areas"
    return None


def _infer_section_from_free_heading(text: str) -> str | None:
    """
    Second pass for Word headings that did not match _match_section_header.
    Uses substring / loose patterns only (short lines).
    """
    t = _normalize_header(text)
    t = _strip_leading_outline(t)
    if len(t) > 90:
        return None
    if "acceptance" in t and (
        "criter" in t or t.strip() in ("acceptance", "acceptance testing")
    ):
        return "acceptance"
    if "success criteria" in t or "exit criteria" in t or "definition of done" in t:
        return "acceptance"
    if re.search(r"\b(requirements?|reqs)\b", t) and len(t) < 75:
        if "test" in t and "case" in t:
            return None
        return "acceptance"
    if "validation" in t and "criter" in t and len(t) < 80:
        return "acceptance"
    if re.search(r"\btest cases?\b", t) and "acceptance" not in t and len(t) < 75:
        return "tests"
    if (
        "test script" in t
        or "test procedure" in t
        or "test steps" in t
        or "test scenarios" in t
        or "uat execution" in t
    ) and len(t) < 80:
        return "tests"
    if re.search(r"\b(uat|qa)\s+test", t) and len(t) < 70:
        return "tests"
    # "TC-01 Login flow" without colon (heading or loose title line)
    if re.match(r"^tc[\s\-]+\d+\s+\S", t) and len(t) < 85:
        remainder = re.sub(r"^tc[\s\-]+\d+\s+", "", t)
        if not re.match(
            r"^(must|should|shall|needs|will\s+be|can\s+be|validates?|ensures?|"
            r"confirms?|passes?|verifies?|is\s+required)\b",
            remainder,
        ):
            return "tests"
    if re.search(r"\b(goal|purpose|objective)\b", t) and len(t) < 70:
        if "acceptance" in t:
            return None
        if re.search(r"\btest\s*(case|plan|script|steps?|scenarios?|execution)\b", t):
            return None
        return "business_goal"
    if "changed" in t and ("area" in t or "file" in t or "list" in t) and len(t) < 85:
        return "changed_areas"
    if ("impacted" in t or "affected" in t) and (
        "module" in t or "component" in t or "system" in t
    ):
        return "changed_areas"
    # "Scope & Context" is workflow narrative, not an impacted-area inventory.
    if re.match(r"^scope\s*(?:&|and)\s*context\b", t) and len(t) < 75:
        return "workflow_ctx"
    if re.search(r"\bscope\b", t) and len(t) < 60 and "test" not in t:
        return "changed_areas"
    if re.search(
        r"\b(integration\s+points?|system\s+dependencies?|external\s+dependencies?)\b",
        t,
    ) and len(t) < 85:
        return "dependencies"
    if re.search(
        r"\b(workflow\s+context|process\s+description|narrative)\b",
        t,
    ) and len(t) < 80:
        return "workflow_ctx"
    if "appendix" in t and "test" in t and len(t) < 85:
        return "tests"
    return None


def _strip_label(text: str, patterns: list[str]) -> str:
    s = text.strip()
    for pat in patterns:
        m = re.match(pat, s, re.IGNORECASE)
        if m:
            return s[m.end() :].strip()
    return s


def _normalize_tc_id(token: str) -> str:
    t = re.sub(r"\s+", "", token.upper())
    m = re.match(r"TC-?(\d+)$", t)
    if m:
        return f"TC-{m.group(1)}"
    return t


def _parse_test_case_header(line: str) -> tuple[str | None, str]:
    """Return (tc_id, title) if line opens a test case."""
    s = line.strip()
    m = re.match(r"^TC\s+(\d+)\s*[:\-–]?\s*(.*)$", s, re.IGNORECASE)
    if m:
        n = int(m.group(1))
        return f"TC-{n:02d}", (m.group(2) or "").strip()
    m = re.match(
        r"^(TC\s*-\s*\d+|TC\d+)\s*[:\-–]?\s*(.*)$", s, re.IGNORECASE
    )
    if m:
        tid = _normalize_tc_id(m.group(1))
        return tid, (m.group(2) or "").strip()
    m = re.match(
        r"^test\s*case\s*(\d+)\s*[:\-–]?\s*(.*)$", s, re.IGNORECASE
    )
    if m:
        return f"TC-{m.group(1)}", (m.group(2) or "").strip()
    m = re.match(
        r"^test\s*case\s*#?\s*(\d+)\s*[:\-–]\s*(.*)$", s, re.IGNORECASE
    )
    if m:
        return f"TC-{m.group(1)}", (m.group(2) or "").strip()
    m = re.match(r"^test\s*case\s*[:\-–]\s*(.+)$", s, re.IGNORECASE)
    if m:
        return None, m.group(1).strip()
    m = re.match(r"^tc\s*:\s*(.+)$", s, re.IGNORECASE)
    if m:
        return None, m.group(1).strip()
    m = re.match(r"^uat\s*[-_]?\s*(\d+)\s*[:\-–]\s*(.*)$", s, re.IGNORECASE)
    if m:
        return f"TC-{m.group(1)}", (m.group(2) or "").strip()
    m = re.match(r"^case\s*(\d+)\s*[:\-–\.]\s*(.*)$", s, re.IGNORECASE)
    if m:
        return f"TC-{m.group(1)}", (m.group(2) or "").strip()
    m = re.match(r"^test\s*#\s*(\d+)\s*[:\-–]?\s*(.*)$", s, re.IGNORECASE)
    if m:
        return f"TC-{m.group(1)}", (m.group(2) or "").strip()
    m = re.match(
        r"^test\s*scenario\s*(\d+)\s*[:\-–\.]\s*(.*)$", s, re.IGNORECASE
    )
    if m:
        return f"TC-{m.group(1)}", (m.group(2) or "").strip()
    m = re.match(r"^scenario\s*(\d+)\s*[:\-–]\s*(.+)$", s, re.IGNORECASE)
    if m and not re.match(r"^scenario\s*:\s*acceptance", s, re.IGNORECASE):
        return f"TC-{m.group(1)}", (m.group(2) or "").strip()
    m = re.match(
        r"^uat\s*case\s*#?\s*(\d+)\s*[:\-–\.]\s*(.*)$", s, re.IGNORECASE
    )
    if m:
        return f"TC-{m.group(1)}", (m.group(2) or "").strip()
    m = re.match(
        r"^validation\s+test\s*#?\s*(\d+)\s*[:\-–\.]\s*(.*)$", s, re.IGNORECASE
    )
    if m:
        n = int(m.group(1))
        return f"TC-{n:02d}", (m.group(2) or "").strip()
    m = re.match(
        r"^test\s*script\s*#?\s*(\d+)\s*[:\-–\.]\s*(.*)$", s, re.IGNORECASE
    )
    if m:
        n = int(m.group(1))
        return f"TC-{n:02d}", (m.group(2) or "").strip()
    return None, ""


def _try_open_test_colon(state: "_DocxParseState", line: str) -> bool:
    """
    Lines like ``Test: Verify checkout`` (not ``Test case:``) open a test case.
    Returns True if this line was consumed as a test opener.
    """
    s = line.strip()
    m = re.match(r"^test\s*[:\-–]\s*(.+)$", s, re.IGNORECASE)
    if not m:
        return False
    title = (m.group(1) or "").strip()
    if not title:
        return False
    if re.match(r"^case\b", title, re.IGNORECASE):
        return False
    state.section = "tests"
    state.open_test_case(None, title)
    return True


def _imperative_step_candidate(s: str) -> bool:
    """Short imperative lines that are obvious UAT actions (not AC/req prose)."""
    t = s.strip()
    if len(t) < 8 or len(t) > 500:
        return False
    if re.match(
        r"^(AC|TC|FR|NFR|BR|REQ|SR|UC)[\s\-]*\d+\b",
        t,
        re.IGNORECASE,
    ):
        return False
    if re.match(r"^(given|when|then)\b", t, re.IGNORECASE):
        return False
    if re.match(
        r"^(the\s+user|users?\s+(?:can|must|should)|the\s+system|system\s+shall)\b",
        t,
        re.IGNORECASE,
    ):
        return False
    if re.match(
        r"^see\s+(figure|fig|screenshot)\b",
        t,
        re.IGNORECASE,
    ):
        return True
    return bool(
        re.match(
            r"^(navigate|click|tap|double[-\s]?click|enter|type|select|choose|open|close|"
            r"submit|scroll|swipe|go\s+to|log\s+(?:in|out)|sign\s+(?:in|out)|wait\s+for|"
            r"press|hit|verify|confirm|validate|ensure|check|observe|record|update|edit|"
            r"delete|remove|add|save|cancel|refresh|load|upload|download|complete)\b",
            t,
            re.IGNORECASE,
        )
    )


def _harvest_structure_from_notes_parts(state: "_DocxParseState") -> None:
    """
    Second pass: promote AC-like / Test:-style lines that landed in notes into structured fields.
    Only moves lines that match explicit patterns (no inference).
    """
    if not state.notes_parts:
        return
    kept: list[str] = []
    for line in state.notes_parts:
        s = (line or "").strip()
        if not s:
            continue
        if re.match(r"^AC[\s\-]*\d+\b", s, re.IGNORECASE):
            state.acceptance_lines.append(_normalize_line_for_parse(s))
            continue
        if re.match(
            r"^(REQ|FR|BR|NFR|SR|UC)\s*[\-\.]?\s*\d+\b",
            s,
            re.IGNORECASE,
        ):
            state.acceptance_lines.append(_normalize_line_for_parse(s))
            continue
        if _numbered_line_prefers_acceptance(s):
            state.acceptance_lines.append(_normalize_line_for_parse(s))
            continue
        if _try_open_test_colon(state, s):
            continue
        tc_id, _title_chk = _parse_test_case_header(s)
        if tc_id is not None:
            state.section = "tests"
            state.open_test_case(tc_id, _title_chk)
            continue
        kept.append(line)
    state.notes_parts = kept


def _docx_structure_trust_warnings(
    acceptance_criteria: list,
    test_cases: list,
    business_goal: str,
    notes_extra: str,
) -> list[str]:
    """Honest extraction-quality hints for sparse or partial DOCX parses."""
    out: list[str] = []
    if not acceptance_criteria:
        out.append(
            "No acceptance criteria extracted; if the document lists AC-# (or FR/REQ ids), place them under an "
            "'Acceptance criteria' heading or keep them out of the Business goal paragraph so they are not merged there."
        )
    if not test_cases:
        out.append(
            "No test cases extracted; use 'Test Case: …', 'TC-01: …', or a line 'Test: …' to open a scenario."
        )
    bg = (business_goal or "").strip()
    if bg and len(bg) > 380 and not bg.startswith("—"):
        out.append(
            "Business goal text is long; confirm it does not include requirement or AC lines that belong under acceptance criteria."
        )
    ne = (notes_extra or "").strip()
    if ne and len(ne) > 120 and (not acceptance_criteria or not test_cases):
        out.append(
            "Notable text remains in notes only; add explicit headings (Acceptance criteria, Test cases) where possible for more reliable parsing."
        )
    return out


def _is_step_line(line: str) -> tuple[bool, str]:
    s = line.strip()
    m = re.match(r"^step\s*(\d+)\s*[:\-–\.]\s*(.+)$", s, re.IGNORECASE)
    if m:
        return True, m.group(2).strip()
    m = re.match(r"^step\s*[:\-–]\s*(.+)$", s, re.IGNORECASE)
    if m and len(m.group(1).strip()) > 3:
        return True, m.group(1).strip()
    m = re.match(r"^(\d+)[\.\)]\s+(.+)$", s)
    if m:
        return True, m.group(2).strip()
    m = re.match(r"^([a-z])[\.\)]\s+(.+)$", s, re.IGNORECASE)
    if m and len(m.group(2).strip()) > 2:
        return True, m.group(2).strip()
    return False, s


def _is_bullet_step_candidate(line: str) -> tuple[bool, str]:
    """Imperative-looking line after '-', '*', or '•' (tests context only)."""
    s = line.strip()
    m = re.match(r"^[\-\*•]\s+(.+)$", s)
    if not m:
        return False, s
    inner = m.group(1).strip()
    if len(inner) < 12 or len(inner) > 500:
        return False, s
    if _match_section_header(inner) or _parse_test_case_header(inner)[0]:
        return False, s
    low = inner[:1].lower() if inner else ""
    if low and inner[0].isdigit():
        return False, s
    return True, inner


def _numbered_line_prefers_acceptance(line: str) -> bool:
    """
    When no test case is open yet, '1. …' lines are often AC bullets, not execution steps.
    Prefer acceptance unless the line reads like an imperative UAT step.
    """
    m = re.match(r"^(\d+)[\.\)]\s+(.+)$", line.strip())
    if not m:
        return False
    body = (m.group(2) or "").strip()
    if not body or len(body) < 3:
        return False
    if re.match(r"^(AC|TC)[\s\-]*\d+\b", body, re.IGNORECASE):
        return False
    if re.match(
        r"^(navigate|click|tap|double[-\s]?click|enter|type|select|choose|open|close|"
        r"submit|scroll|swipe|go\s+to|log\s+(?:in|out)|sign\s+(?:in|out)|wait\s+for|"
        r"press|hit)\b",
        body,
        re.IGNORECASE,
    ):
        return False
    if re.match(r"^(given|when|then)\b", body, re.IGNORECASE):
        return True
    if re.search(
        r"\b(shall|must\s+be|should\s+be|system\s+shall|user\s+(?:can|must|should)|"
        r"is\s+able\s+to|will\s+(?:allow|display|show|validate)|verify\s+that|"
        r"ensures?|allows?|enables?|successful\s+(?:login|authentication|response))\b",
        body,
        re.IGNORECASE,
    ):
        return True
    if re.match(
        r"^(the\s+user|users?\s+(?:can|must|should)|the\s+system|a\s+registered\s+user)\b",
        body,
        re.IGNORECASE,
    ):
        return True
    return False


def _bullet_body_prefers_acceptance(body: str) -> bool:
    """Same idea as _numbered_line_prefers_acceptance for '- …' / '• …' lines."""
    b = (body or "").strip()
    if len(b) < 12:
        return False
    if _match_section_header(b) or _parse_test_case_header(b)[0]:
        return False
    return _numbered_line_prefers_acceptance(f"1. {b}")


def _is_dependency_table_header_pair(c0: str, c1: str) -> bool:
    a = _normalize_header(c0)
    b = _normalize_header(c1)
    if not a or not b:
        return False
    left_ok = bool(
        re.match(
            r"^(dependency|dependencies|system|interface|integration)(\s+name|\s+id)?$",
            a,
        )
    )
    right_ok = bool(
        re.match(r"^(description|details|notes|owner|impact)$", b)
    )
    return left_ok and right_ok


def _is_changed_area_table_header_pair(c0: str, c1: str) -> bool:
    a = _normalize_header(c0)
    b = _normalize_header(c1)
    if not a or not b:
        return False
    left_ok = bool(
        re.match(
            r"^(application|module|area|component|screen|subsystem|product)$",
            a,
        )
    )
    right_ok = bool(
        re.match(r"^(description|change|summary|type|impact|details)$", b)
    )
    return left_ok and right_ok


def _is_ac_table_header_pair(c0: str, c1: str) -> bool:
    a = _normalize_header(c0)
    b = _normalize_header(c1)
    if not a or not b:
        return False
    left_ok = bool(
        re.match(
            r"^(ac(\s*#|\s*id)?|req(\.\s*)?id|requirement(\s*id)?|requirement|id|#|story\s*id)$",
            a,
        )
    )
    right_ok = bool(
        re.match(
            r"^(description|details|criteria|acceptance|text|summary|statement|user\s*story)$",
            b,
        )
    )
    return left_ok and right_ok


# Figure / screenshot labels in captions (Figure 1, Fig 2, Figure 1.1, [Screenshot 3])
_FIG_LABEL_INLINE = re.compile(
    r"(?i)\b(?:figure|fig\.?)\s*([\d]+(?:\.[\d]+)*)\b",
)
_BRACKET_SCREEN_LABEL = re.compile(
    r"(?i)\[\s*(?:screenshot|screen|fig\.?|figure)\s*([\d]+(?:\.[\d]+)*)\s*\]",
)

# References inside step prose ("see Figure 1.1"); contextual patterns first
_FIG_REF_CONTEXT = re.compile(
    r"(?i)(?:see|refer\s+to|as\s+shown\s+in|according\s+to|per|in|from)\s+(?:the\s+)?"
    r"(?:figure|fig\.?)\s*([\d]+(?:\.[\d]+)*)",
)


def _normalize_fig_key(key: str) -> str:
    key = (key or "").strip()
    if not key:
        return key
    parts = key.split(".")
    out: list[str] = []
    for p in parts:
        p = p.strip()
        if p.isdigit():
            out.append(str(int(p)))
        else:
            out.append(p)
    return ".".join(out)


def _extract_figure_labels_from_text(text: str) -> list[str]:
    if not text or not text.strip():
        return []
    keys: list[str] = []
    seen: set[str] = set()
    for rx in (_FIG_LABEL_INLINE, _BRACKET_SCREEN_LABEL):
        for m in rx.finditer(text):
            k = _normalize_fig_key(m.group(1))
            if k and k not in seen:
                seen.add(k)
                keys.append(k)
    return keys


def _extract_figure_refs_from_step(text: str) -> list[str]:
    """Ordered unique figure keys referenced in step text."""
    if not text or not text.strip():
        return []
    keys: list[str] = []
    seen: set[str] = set()
    for m in _FIG_REF_CONTEXT.finditer(text):
        k = _normalize_fig_key(m.group(1))
        if k and k not in seen:
            seen.add(k)
            keys.append(k)
    if not keys:
        for m in _FIG_LABEL_INLINE.finditer(text):
            k = _normalize_fig_key(m.group(1))
            if k and k not in seen:
                seen.add(k)
                keys.append(k)
    return keys


def _looks_like_standalone_caption(text: str) -> bool:
    """True when paragraph is short and not a numbered step line."""
    t = text.strip()
    if not t or len(t) > 220:
        return False
    if _is_step_line(t)[0]:
        return False
    return bool(_extract_figure_labels_from_text(t))


def _step_item_to_str(step_item: object) -> str:
    if isinstance(step_item, str):
        return step_item
    if isinstance(step_item, dict):
        return str(step_item.get("text") or "")
    return str(step_item or "")


def _register_figure_labels(
    state: "_DocxParseState",
    paths: list[str],
    raw_text: str,
) -> None:
    """Map caption labels to saved image paths (last write wins on duplicate keys)."""
    labels = _extract_figure_labels_from_text(raw_text)
    if paths:
        if state.pending_figure_label:
            pk = _normalize_fig_key(state.pending_figure_label)
            for p in paths:
                if pk:
                    state.figure_registry[pk] = p
            state.pending_figure_label = None
        if labels:
            if len(labels) == 1:
                for p in paths:
                    state.figure_registry[labels[0]] = p
            else:
                for i, p in enumerate(paths):
                    if i < len(labels):
                        state.figure_registry[labels[i]] = p
        return
    if labels and _looks_like_standalone_caption(raw_text):
        state.pending_figure_label = labels[0]


def _compact_parallel_screenshots_to_objects(
    paths: list[str], labels: list[str]
) -> list[dict]:
    """
    Drop empty per-step slots; keep schema shape via mapped_to_step_index objects only.
    """
    out: list[dict] = []
    for i, path in enumerate(paths):
        ps = (path or "").strip()
        if not ps:
            continue
        entry: dict = {"path": ps, "mapped_to_step_index": i}
        lb = labels[i] if i < len(labels) else ""
        if lb and str(lb).strip():
            entry["label"] = str(lb).strip()
        out.append(entry)
    return out


def _set_step_item_text(tc: dict, index: int, new_text: str) -> None:
    steps = tc.get("steps") or []
    if index < 0 or index >= len(steps):
        return
    it = steps[index]
    if isinstance(it, dict):
        it["text"] = new_text
    else:
        steps[index] = new_text


def _strip_mapped_figure_prose_from_step(text: str, refs: list[str]) -> str:
    """
    Remove contextual 'see Figure N' clauses only for refs that were explicitly mapped
    to this step. If nothing remains, returns original text (caller skips update).
    """
    if not text or not text.strip() or not refs:
        return text
    original = text
    t = text
    for ref in refs:
        rk = (ref or "").strip()
        if not rk:
            continue
        esc = re.escape(rk)
        clause = (
            r"(?:\s*[.,;]?\s*)"
            r"(?:see|refer\s+to|as\s+shown\s+in|according\s+to|per|in|from)\s+"
            r"(?:the\s+)?(?:figure|fig\.?)\s*" + esc + r"\b\.?"
        )
        t = re.sub(clause, "", t, flags=re.IGNORECASE)
    t = re.sub(r"\s{2,}", " ", t).strip(" \t,;.-")
    t = t.strip()
    if not t:
        return original
    return t


def _collect_placed_media_paths(test_cases: list[dict], workflow: list[str]) -> set[str]:
    out: set[str] = set()
    for tc in test_cases:
        for item in tc.get("expected_step_screenshots") or []:
            if isinstance(item, str):
                s = item.strip()
                if s:
                    out.add(s)
            elif isinstance(item, dict):
                p = item.get("path")
                if p and str(p).strip():
                    out.add(str(p).strip())
    for p in workflow:
        if isinstance(p, str) and p.strip():
            out.add(p.strip())
    return out


def _registry_figure_keys_for_placed_paths(
    registry: dict[str, str], placed_paths: set[str]
) -> set[str]:
    keys: set[str] = set()
    for k, p in registry.items():
        if p and str(p).strip() in placed_paths:
            keys.add(_normalize_fig_key(k))
    return keys


def _filter_redundant_figure_caption_lines(
    notes_extra: str, placed_fig_keys: set[str]
) -> str:
    if not (notes_extra or "").strip() or not placed_fig_keys:
        return notes_extra
    lines = notes_extra.splitlines()
    kept: list[str] = []
    for line in lines:
        s = line.strip()
        if not s:
            kept.append(line)
            continue
        keys = [_normalize_fig_key(x) for x in _extract_figure_labels_from_text(s)]
        if (
            keys
            and _looks_like_standalone_caption(s)
            and all(k in placed_fig_keys for k in keys)
        ):
            continue
        kept.append(line)
    return "\n".join(kept).strip()


def _reconcile_figure_mappings_for_test_cases(
    test_cases: list[dict],
    registry: dict[str, str],
    workflow_sink: list[str],
) -> None:
    """
    Rebuild expected_step_screenshots per test case: explicit figure refs in steps
    override order-based placement; remaining heuristic paths fill empty slots;
    leftovers go to workflow_sink.
    """
    if not registry:
        for tc in test_cases:
            p, lb = resolved_step_screenshots(tc)
            compact = _compact_parallel_screenshots_to_objects(p, lb)
            tc["expected_step_screenshots"] = compact
            tc.pop("expected_step_screenshot_labels", None)
        return

    for tc in test_cases:
        steps_raw = tc.get("steps") or []
        n = len(steps_raw)
        if n == 0:
            continue
        heur = raw_step_screenshot_paths_in_json_order(tc)
        paths_out = [""] * n
        labels_out = [""] * n
        used_paths: set[str] = set()
        explicit_step: set[int] = set()

        for i in range(n):
            step_text = _step_item_to_str(steps_raw[i])
            for ref in _extract_figure_refs_from_step(step_text):
                p = registry.get(ref)
                if p and p not in used_paths:
                    paths_out[i] = p
                    labels_out[i] = ref
                    used_paths.add(p)
                    explicit_step.add(i)
                    break

        hi = 0
        for i in range(n):
            if paths_out[i]:
                continue
            while hi < len(heur) and (not heur[hi] or heur[hi] in used_paths):
                hi += 1
            if hi < len(heur):
                paths_out[i] = heur[hi]
                used_paths.add(heur[hi])
                hi += 1

        for p in heur:
            if p and p not in used_paths:
                workflow_sink.append(p)

        for i in sorted(explicit_step):
            ref = labels_out[i] if i < len(labels_out) else ""
            if not ref:
                continue
            raw_text = _step_item_to_str(steps_raw[i])
            cleaned = _strip_mapped_figure_prose_from_step(raw_text, [ref])
            if cleaned != raw_text:
                _set_step_item_text(tc, i, cleaned)

        compact = _compact_parallel_screenshots_to_objects(paths_out, labels_out)
        tc["expected_step_screenshots"] = compact
        tc.pop("expected_step_screenshot_labels", None)


def _append_registry_orphans_to_workflow(
    registry: dict[str, str],
    test_cases: list[dict],
    workflow_sink: list[str],
) -> None:
    placed: set[str] = set()
    for tc in test_cases:
        for p in resolved_step_screenshots(tc)[0]:
            if p:
                placed.add(str(p))
    for p in registry.values():
        if p and p not in placed:
            workflow_sink.append(p)


class _DocxParseState:
    def __init__(self) -> None:
        self.section: str | None = None
        self.story_title = ""
        self.story_description_parts: list[str] = []
        self.business_goal_parts: list[str] = []
        self.acceptance_lines: list[str] = []
        self.changed_area_lines: list[str] = []
        self.dependency_lines: list[str] = []
        self.workflow_context_parts: list[str] = []
        self.notes_parts: list[str] = []
        self.test_cases: list[dict] = []
        self.workflow_images: list[str] = []
        self._buf: list[str] = []
        self.figure_registry: dict[str, str] = {}
        self.pending_figure_label: str | None = None

    def flush_description_buffer(self, target: str) -> None:
        if not self._buf:
            return
        text = "\n".join(self._buf).strip()
        self._buf = []
        if not text:
            return
        if target == "story_description":
            self.story_description_parts.append(text)
        elif target == "business_goal":
            self.business_goal_parts.append(text)
        elif target == "acceptance":
            self.acceptance_lines.extend(text.splitlines())
        elif target == "notes":
            self.notes_parts.append(text)

    def open_test_case(self, tc_id: str | None, title: str) -> None:
        tid = tc_id or f"TC-{len(self.test_cases) + 1:02d}"
        seen = {t["id"] for t in self.test_cases}
        if tid in seen:
            n = 2
            while f"{tid}-dup{n}" in seen:
                n += 1
            tid = f"{tid}-dup{n}"
        self.test_cases.append(
            {
                "id": tid,
                "text": title or f"Test case {len(self.test_cases) + 1}",
                "steps": [],
                "expected_step_screenshots": [],
            }
        )

    def add_step(self, text: str) -> None:
        if not self.test_cases:
            self.open_test_case(None, "")
        self.test_cases[-1]["steps"].append(text)

    def assign_images(self, paths: list[str]) -> None:
        if not paths:
            return
        # Attach to step slots only in active test execution context; otherwise workflow
        # / dependency / acceptance / workflow_ctx images stay in workflow_process_screenshots
        # (avoids mis-mapping healthcare "context" screenshots onto unrelated TCs).
        in_tests_context = self.section == "tests" or (
            self.section is None
            and self.test_cases
            and len(self.test_cases[-1].get("steps") or []) > 0
        )
        if not self.test_cases or not in_tests_context:
            self.workflow_images.extend(paths)
            return
        tc = self.test_cases[-1]
        steps: list = tc.get("steps") or []
        shots: list = tc.setdefault("expected_step_screenshots", [])
        for rel in paths:
            if len(steps) == 0:
                self.workflow_images.append(rel)
            elif len(shots) < len(steps):
                shots.append(rel)
            else:
                # All step slots already have an image; avoid attaching to the wrong step.
                self.workflow_images.append(rel)


def _process_paragraph(state: _DocxParseState, paragraph: Paragraph) -> None:
    """Dispatch body vs heading paragraphs; preserve multi-line bodies."""
    raw_full = (paragraph.text or "").strip()
    if _is_heading_paragraph(paragraph) and raw_full:
        _process_line(state, raw_full, paragraph=None, heading=True)
        return
    for raw_line in (paragraph.text or "").splitlines():
        line = _normalize_line_for_parse(raw_line)
        if line:
            _process_line(state, line, paragraph=paragraph, heading=False)


def _dedupe_preserve_order(items: list[str]) -> list[str]:
    seen: set[str] = set()
    out: list[str] = []
    for x in items:
        k = x.strip()
        if k and k not in seen:
            seen.add(k)
            out.append(k)
    return out


def _table_row_texts(row) -> list[str]:
    texts: list[str] = []
    for cell in row.cells:
        parts: list[str] = []
        for p in cell.paragraphs:
            t = p.text.strip()
            if t:
                parts.append(t)
        merged_cell = " ".join(parts).strip()
        if merged_cell:
            texts.append(merged_cell)
    return _dedupe_preserve_order(texts)


def _emit_table_row(state: _DocxParseState, cell_texts: list[str]) -> None:
    if not cell_texts:
        return
    ct0 = (cell_texts[0] or "").strip()
    c1 = (cell_texts[1] or "").strip() if len(cell_texts) >= 2 else ""
    if len(cell_texts) >= 2 and _is_ac_table_header_pair(ct0, c1):
        logger.debug("docx table: skipped AC/requirement header row")
        return
    if len(cell_texts) >= 2 and _is_dependency_table_header_pair(ct0, c1):
        logger.debug("docx table: skipped dependency header row")
        return
    if len(cell_texts) >= 2 and _is_changed_area_table_header_pair(ct0, c1):
        logger.debug("docx table: skipped changed-area header row")
        return
    ac_id_row = re.match(
        r"^(AC[\s\-]*\d+|(?:FR|NFR|BR|SR|REQ|UC)[\s\-\.]?\d+)\b",
        ct0,
        re.IGNORECASE,
    )
    if len(cell_texts) >= 2 and ac_id_row:
        if state.section not in ("tests",):
            state.section = "acceptance"
        ac_line = f"{ct0} — {' | '.join(cell_texts[1:])}"
        state.acceptance_lines.append(_normalize_line_for_parse(ac_line))
        logger.debug("docx table: AC/requirement row -> acceptance_lines")
        return
    if len(cell_texts) >= 2 and state.section in (
        None,
        "dependencies",
        "changed_areas",
    ):
        tnorm = _normalize_header(ct0)
        # Prefer app/module/screen rows as changed areas (even mid dependency tables).
        if re.match(
            r"^(application|module|component|screen|subsystem|area|product|patient|"
            r"clinical|record|portal|billing|encounter|order)\b",
            tnorm,
        ):
            if state.section in (None, "dependencies"):
                state.section = "changed_areas"
            row = f"{ct0} — {c1}".strip()
            if _is_changed_areas_source_line_junk(row) or _is_junk_changed_area_row(
                {"area": ct0.strip(), "type": c1.strip()}
            ):
                logger.debug("docx table: skipped changed-area junk row")
                return
            state.changed_area_lines.append(_normalize_line_for_parse(row))
            logger.debug("docx table: changed-area-style row")
            return
        if state.section in (None, "dependencies") and (
            re.match(
                r"^(system|interface|integration|dependency|vendor)\b",
                tnorm,
            )
            or (
                state.section == "dependencies"
                and len(ct0) <= 48
                and len(c1) >= 8
                and not re.match(
                    r"^(action|step|expected|result|test|application|module)\b",
                    tnorm,
                )
            )
        ):
            if re.match(
                r"^(system|interface|integration|dependency|vendor)\b",
                tnorm,
            ) or (
                state.section == "dependencies"
                and not re.match(
                    r"^(action|step|expected|result|test)\b",
                    tnorm,
                )
            ):
                if state.section is None:
                    state.section = "dependencies"
                row = f"{ct0} — {c1}".strip()
                state.dependency_lines.append(_normalize_line_for_parse(row))
                logger.debug("docx table: dependency-style row")
                return
    if len(cell_texts) >= 2:
        tc_id, title0 = _parse_test_case_header(ct0)
        if tc_id is not None:
            extra = " | ".join(cell_texts[1:]).strip()
            title = f"{title0} {extra}".strip() if title0 else extra
            merged = f"{tc_id}: {title}".strip() if title else f"{tc_id}:"
            _process_line(state, merged, paragraph=None, heading=False)
            return
    if len(cell_texts) >= 2 and state.section == "tests":
        merged = f"{cell_texts[0]} — Expected: {' | '.join(cell_texts[1:])}"
    else:
        merged = " | ".join(cell_texts)
    _process_line(state, merged, paragraph=None, heading=False)


def _process_line(
    state: _DocxParseState,
    line: str,
    paragraph: Paragraph | None = None,
    *,
    heading: bool = False,
) -> None:
    if not line.strip():
        return

    hdr = _match_section_header(line)
    if heading and hdr is None:
        hdr = _infer_section_from_free_heading(line)
    if heading and hdr is None:
        logger.debug("docx unmatched heading -> notes: %r", line[:120])
        state.notes_parts.append(line.strip())
        return

    if hdr:
        if state.section == "story_description" and hdr != "story_description":
            state.flush_description_buffer("story_description")
        if state.section == "business_goal" and hdr != "business_goal":
            state.flush_description_buffer("business_goal")
        if state.section == "acceptance" and hdr not in ("acceptance",):
            state.flush_description_buffer("acceptance")

        state.section = hdr
        logger.debug("docx section header -> %s (line=%r)", hdr, line[:120])
        if hdr != "tests":
            state.pending_figure_label = None
        if hdr == "story_title":
            rest = _strip_label(
                line,
                [
                    r"^(story\s*)?title\s*:",
                    r"^title\s*:",
                ],
            )
            if rest:
                state.story_title = rest
        elif hdr == "story_description":
            rest = _strip_label(
                line,
                [
                    r"^(story\s*)?description\s*:",
                    r"^overview\s*:",
                    r"^user\s*story\s*:",
                    r"^scenario\s*:",
                ],
            )
            state._buf = [rest] if rest else []
        elif hdr == "business_goal":
            rest = _strip_label(
                line,
                [
                    r"^business\s*goals?\s*:",
                    r"^business\s*goal\s*:",
                    r"^business\s*objectives?\s*:",
                    r"^objectives?\s*:",
                    r"^objective\s*:",
                    r"^goal\s*:",
                    r"^goals\s*:",
                    r"^purpose\s*:",
                    r"^strategic\s*(goal|objective)s?\s*:",
                ],
            )
            bare = _normalize_header(rest)
            if bare in (
                "purpose",
                "goal",
                "goals",
                "objective",
                "objectives",
                "business goal",
                "business goals",
                "business objective",
                "business objectives",
            ):
                rest = ""
            state._buf = [rest] if rest else []
        elif hdr == "acceptance":
            rest = _strip_label(
                line,
                [
                    r"^acceptance\s*criteria\s*:?",
                    r"^requirements\s*:?",
                    r"^functional\s*requirements?\s*:?",
                    r"^product\s*requirements?\s*:?",
                    r"^business\s*requirements?\s*:?",
                    r"^user\s*requirements?\s*:?",
                    r"^success\s*criteria\s*:?",
                    r"^validation\s*criteria\s*:?",
                    r"^exit\s*criteria\s*:?",
                ],
            )
            if _normalize_header(rest) in (
                "requirements",
                "requirement",
                "acceptance criteria",
                "acceptance",
                "functional requirements",
                "user requirements",
                "product requirements",
                "business requirements",
            ):
                rest = ""
            if rest:
                state.acceptance_lines.append(_normalize_line_for_parse(rest))
            state._buf = []
        elif hdr == "changed_areas":
            rest = _strip_label(
                line,
                [
                    r"^changed\s*areas?\s*:?",
                    r"^scope\s*of\s*change\s*:?",
                    r"^impacted\s*areas?\s*:?",
                    r"^affected\s*areas?\s*:?",
                    r"^scope\s*&\s*context\s*:?",
                    r"^scope\s*&\s*context$",
                    r"^scope\s*:?",
                    r"^systems?\s*impacted\s*:?",
                    r"^application\s*impact\s*:?",
                ],
            )
            if rest:
                state.changed_area_lines.append(_normalize_line_for_parse(rest))
            state._buf = []
        elif hdr == "dependencies":
            rest = _strip_label(
                line,
                [
                    r"^known\s*dependencies?\s*:?",
                    r"^dependencies?\s*:?",
                    r"^prerequisites?\s*:?",
                    r"^system\s*dependencies?\s*:?",
                    r"^integration\s*(points?|interfaces?)?\s*:?",
                    r"^external\s*systems?\s*:?",
                    r"^interfaces?\s*:?",
                ],
            )
            if rest:
                state.dependency_lines.append(_normalize_line_for_parse(rest))
            state._buf = []
        elif hdr == "tests":
            state._buf = []
            tc_id, title = _parse_test_case_header(line)
            if tc_id is not None:
                state.open_test_case(tc_id, title)
            elif title and re.match(r"^test\s*case\b", line, re.IGNORECASE):
                state.open_test_case(None, title)
        elif hdr == "workflow_ctx":
            state._buf = []
        return

    # Continuation of multi-line section fields
    if state.section == "story_title":
        hdr = _match_section_header(line)
        if hdr:
            state.section = None
            _process_line(state, line, paragraph=paragraph, heading=False)
            return
        tc_id, title = _parse_test_case_header(line)
        if tc_id is not None:
            state.section = "tests"
            state.open_test_case(tc_id, title)
            return
        if title and re.match(r"^test\s*case\b", line, re.IGNORECASE):
            state.section = "tests"
            state.open_test_case(None, title)
            return
        if re.match(r"^AC[\s\-]*\d+\b", line, re.IGNORECASE) or (
            _numbered_line_prefers_acceptance(line)
        ):
            state.section = None
            _process_line(state, line, paragraph=paragraph, heading=False)
            return
        bs, btext = _is_bullet_step_candidate(line)
        if bs and _bullet_body_prefers_acceptance(btext):
            state.section = None
            _process_line(state, line, paragraph=paragraph, heading=False)
            return
        rest = line.strip()
        if rest:
            state.story_title = (
                f"{state.story_title} {rest}".strip() if state.story_title else rest
            )
        return
    if state.section == "story_description":
        hdr = _match_section_header(line)
        if hdr:
            state.flush_description_buffer("story_description")
            state.section = None
            _process_line(state, line, paragraph=paragraph, heading=False)
            return
        tc_id, title = _parse_test_case_header(line)
        if tc_id is not None:
            state.flush_description_buffer("story_description")
            state.section = "tests"
            state.open_test_case(tc_id, title)
            return
        if title and re.match(r"^test\s*case\b", line, re.IGNORECASE):
            state.flush_description_buffer("story_description")
            state.section = "tests"
            state.open_test_case(None, title)
            return
        if _try_open_test_colon(state, line):
            state.flush_description_buffer("story_description")
            return
        if re.match(r"^AC[\s\-]*\d+\b", line, re.IGNORECASE) or (
            _numbered_line_prefers_acceptance(line)
        ):
            state.flush_description_buffer("story_description")
            state.section = None
            _process_line(state, line, paragraph=paragraph, heading=False)
            return
        bs, btext = _is_bullet_step_candidate(line)
        if bs and _bullet_body_prefers_acceptance(btext):
            state.flush_description_buffer("story_description")
            state.section = None
            _process_line(state, line, paragraph=paragraph, heading=False)
            return
        state._buf.append(line)
        return
    if state.section == "business_goal":
        hdr = _match_section_header(line)
        if hdr:
            state.flush_description_buffer("business_goal")
            state.section = None
            _process_line(state, line, paragraph=paragraph, heading=False)
            return
        tc_id, title = _parse_test_case_header(line)
        if tc_id is not None:
            state.flush_description_buffer("business_goal")
            state.section = "tests"
            state.open_test_case(tc_id, title)
            return
        if title and re.match(r"^test\s*case\b", line, re.IGNORECASE):
            state.flush_description_buffer("business_goal")
            state.section = "tests"
            state.open_test_case(None, title)
            return
        if _try_open_test_colon(state, line):
            state.flush_description_buffer("business_goal")
            return
        if re.match(r"^AC[\s\-]*\d+\b", line, re.IGNORECASE):
            state.flush_description_buffer("business_goal")
            state.section = None
            _process_line(state, line, paragraph=paragraph, heading=False)
            return
        if _numbered_line_prefers_acceptance(line):
            state.flush_description_buffer("business_goal")
            state.section = None
            _process_line(state, line, paragraph=paragraph, heading=False)
            return
        bs, btext = _is_bullet_step_candidate(line)
        if bs and _bullet_body_prefers_acceptance(btext):
            state.flush_description_buffer("business_goal")
            state.section = None
            _process_line(state, line, paragraph=paragraph, heading=False)
            return
        if re.match(
            r"^(REQ|FR|BR|NFR|SR|UC)\s*[\-\.]?\s*\d+\b",
            line,
            re.IGNORECASE,
        ):
            state.flush_description_buffer("business_goal")
            state.section = None
            _process_line(state, line, paragraph=paragraph, heading=False)
            return
        state._buf.append(line)
        return
    if state.section == "workflow_ctx":
        hdr = _match_section_header(line)
        if hdr:
            state.section = None
            _process_line(state, line, paragraph=paragraph, heading=False)
            return
        tc_id, title = _parse_test_case_header(line)
        if tc_id is not None:
            state.section = "tests"
            state.open_test_case(tc_id, title)
            return
        if title and re.match(r"^test\s*case\b", line, re.IGNORECASE):
            state.section = "tests"
            state.open_test_case(None, title)
            return
        if _try_open_test_colon(state, line):
            return
        if re.match(r"^AC[\s\-]*\d+\b", line, re.IGNORECASE):
            state.section = None
            _process_line(state, line, paragraph=paragraph, heading=False)
            return
        if _numbered_line_prefers_acceptance(line):
            state.section = None
            _process_line(state, line, paragraph=paragraph, heading=False)
            return
        bs, btext = _is_bullet_step_candidate(line)
        if bs and _bullet_body_prefers_acceptance(btext):
            state.section = None
            _process_line(state, line, paragraph=paragraph, heading=False)
            return
        if re.match(
            r"^(REQ|FR|BR|NFR|SR|UC)\s*[\-\.]?\s*\d+\b",
            line,
            re.IGNORECASE,
        ):
            state.section = None
            _process_line(state, line, paragraph=paragraph, heading=False)
            return
        state.workflow_context_parts.append(line)
        return
    if state.section == "acceptance":
        hdr2 = _match_section_header(line)
        if hdr2:
            state.section = None
            _process_line(state, line, paragraph=paragraph, heading=False)
            return
        tc_id, title = _parse_test_case_header(line)
        if tc_id is not None:
            state.section = "tests"
            state.open_test_case(tc_id, title)
            return
        if title and re.match(r"^test\s*case\b", line, re.IGNORECASE):
            state.section = "tests"
            state.open_test_case(None, title)
            return
        if _try_open_test_colon(state, line):
            state.flush_description_buffer("acceptance")
            return
        state.acceptance_lines.append(_normalize_line_for_parse(line))
        return
    if state.section == "changed_areas":
        hdr2 = _match_section_header(line)
        if hdr2:
            state.section = None
            _process_line(state, line, paragraph=paragraph, heading=False)
            return
        tc_id, title = _parse_test_case_header(line)
        if tc_id is not None:
            state.section = "tests"
            state.open_test_case(tc_id, title)
            return
        if title and re.match(r"^test\s*case\b", line, re.IGNORECASE):
            state.section = "tests"
            state.open_test_case(None, title)
            return
        state.changed_area_lines.append(_normalize_line_for_parse(line))
        return
    if state.section == "dependencies":
        hdr2 = _match_section_header(line)
        if hdr2:
            state.section = None
            _process_line(state, line, paragraph=paragraph, heading=False)
            return
        tc_id, title = _parse_test_case_header(line)
        if tc_id is not None:
            state.section = "tests"
            state.open_test_case(tc_id, title)
            return
        if title and re.match(r"^test\s*case\b", line, re.IGNORECASE):
            state.section = "tests"
            state.open_test_case(None, title)
            return
        state.dependency_lines.append(_normalize_line_for_parse(line))
        return

    if state.section == "tests" or state.section is None:
        if state.section == "tests" and state.test_cases and not heading:
            stripped = line.strip()
            if re.match(r"^steps?\s*[:\-–]\s*$", stripped, re.IGNORECASE):
                return
            sm = re.match(r"^steps?\s*[:\-–]\s*(.+)$", stripped, re.IGNORECASE)
            if sm and sm.group(1).strip():
                state.add_step(sm.group(1).strip())
                return
            if _imperative_step_candidate(stripped):
                state.add_step(stripped)
                return
        if state.section is None and not state.test_cases and not heading:
            bs, btext = _is_bullet_step_candidate(line)
            if bs and _bullet_body_prefers_acceptance(btext):
                state.section = "acceptance"
                state.acceptance_lines.append(_normalize_line_for_parse(btext))
                logger.debug("docx unscoped bullet line -> acceptance")
                return
            if _numbered_line_prefers_acceptance(line):
                state.section = "acceptance"
                state.acceptance_lines.append(_normalize_line_for_parse(line))
                logger.debug("docx unscoped numbered line -> acceptance")
                return
        if state.section is None and re.match(
            r"^(REQ|FR|BR|NFR|SR|UC)\s*[\-\.]?\s*\d+\b",
            line,
            re.IGNORECASE,
        ):
            state.section = "acceptance"
            state.acceptance_lines.append(_normalize_line_for_parse(line))
            return
        if state.section is None and len(line) < 300 and re.match(
            r"^(Given|When|Then)\s+",
            line,
            re.IGNORECASE,
        ):
            state.section = "acceptance"
            state.acceptance_lines.append(_normalize_line_for_parse(line))
            return
        if state.section is None and re.match(
            r"^(Module|Component|Screen|Page|Service|API|Microservice|Database)\s*:\s*\S",
            line,
            re.IGNORECASE,
        ):
            state.section = "changed_areas"
            state.changed_area_lines.append(_normalize_line_for_parse(line))
            return
        if state.section is None and re.match(
            r"^(scope|in[-\s]?scope|out[-\s]?of[-\s]?scope)\s*:",
            line,
            re.IGNORECASE,
        ):
            state.section = "changed_areas"
            rest = _strip_label(
                line,
                [
                    r"^scope\s*:",
                    r"^in[-\s]?scope\s*:",
                    r"^out[-\s]?of[-\s]?scope\s*:",
                ],
            )
            if rest:
                state.changed_area_lines.append(_normalize_line_for_parse(rest))
            logger.debug("docx unscoped scope line -> changed_areas")
            return
        if state.section is None and re.match(
            r"^(system|interface|integration|external\s+system)\s*[:\-–]\s*\S",
            line,
            re.IGNORECASE,
        ):
            state.section = "dependencies"
            state.dependency_lines.append(_normalize_line_for_parse(line))
            logger.debug("docx unscoped labeled dependency line")
            return
        if state.section is None and re.match(r"^AC[\s-]*\d+\b", line, re.IGNORECASE):
            state.section = "acceptance"
            _process_line(state, line, paragraph=paragraph, heading=False)
            return
        if state.section in (None, "tests") and _try_open_test_colon(state, line):
            return
        if _is_tests_subsection_label(line):
            return
        erm = re.match(
            r"^(expected|actual)\s*(result|results?|outcome|output)?\s*[:\-–]\s*(.+)$",
            line,
            re.IGNORECASE,
        )
        if erm and _merge_expected_result_into_last_step(state, erm.group(3)):
            state.section = "tests"
            return
        erm_short = re.match(r"^expected\s*[:\-–]\s*(.+)$", line, re.IGNORECASE)
        if erm_short and _merge_expected_result_into_last_step(
            state, erm_short.group(1)
        ):
            state.section = "tests"
            return
        tc_id, title = _parse_test_case_header(line)
        if tc_id is not None:
            state.section = "tests"
            state.open_test_case(tc_id, title)
            return
        if title and re.match(r"^test\s*case\b", line, re.IGNORECASE):
            state.section = "tests"
            state.open_test_case(None, title)
            return
        if not heading:
            if " — Expected: " in line and (
                state.section == "tests" or state.test_cases
            ):
                state.section = "tests"
                state.add_step(line.strip())
                return
            is_step, step_text = _is_step_line(line)
            if is_step:
                state.section = "tests"
                state.add_step(step_text)
                return
            if (state.section == "tests" or state.test_cases) and not heading:
                bs, btext = _is_bullet_step_candidate(line)
                if bs:
                    state.section = "tests"
                    state.add_step(btext)
                    return
            if (
                paragraph is not None
                and _paragraph_has_numbering(paragraph)
                and line.strip()
                and _match_section_header(line) is None
                and (state.section == "tests" or state.test_cases)
            ):
                state.section = "tests"
                state.add_step(line.strip())
                return
        if state.section is None and state.test_cases and not heading:
            stripped = line.strip()
            if re.match(r"^steps?\s*[:\-–]\s*$", stripped, re.IGNORECASE):
                return
            sm = re.match(r"^steps?\s*[:\-–]\s*(.+)$", stripped, re.IGNORECASE)
            if sm and sm.group(1).strip():
                state.section = "tests"
                state.add_step(sm.group(1).strip())
                return
            if _imperative_step_candidate(stripped):
                state.section = "tests"
                state.add_step(stripped)
                return
        if state.section == "tests" and state.test_cases:
            state.notes_parts.append(line)
            return
        state.notes_parts.append(line)
        return

    state.notes_parts.append(line)


def _docx_compare_normalize(s: str) -> str:
    """Lowercase + collapse space; normalize common Unicode ampersand/spaces for matching."""
    if not s:
        return ""
    t = str(s).replace("\uff06", "&").replace("\ufeff", "")
    t = t.replace("\u00a0", " ").replace("\u2009", " ")
    return re.sub(r"\s+", " ", t).strip().lower()


def _is_junk_changed_area_row(row: dict) -> bool:
    """Drop subsection bleed / navigation lines mistaken for changed areas."""
    area_raw = (row.get("area") or "").strip()
    typ_raw = (row.get("type") or "").strip()
    if not area_raw and not typ_raw:
        return True
    al = _docx_compare_normalize(area_raw)
    tl = _docx_compare_normalize(typ_raw)
    if re.fullmatch(r"&\s*context\.?", al) or re.fullmatch(r"&\s*context\.?", tl):
        return True
    if al in ("& context", "context") and len(area_raw) <= 28:
        return True
    if tl in ("& context", "context") and len(typ_raw) <= 28:
        return True
    if re.fullmatch(r"[\s\-–—(]*&\s*context\.?", al) or re.fullmatch(
        r"[\s\-–—(]*&\s*context\.?", tl
    ):
        return True
    combined = f"{al} {tl}".strip()
    if combined in ("scope & context", "scope&context"):
        return True
    if al == "scope" and re.fullmatch(r"&?\s*context\.?", tl):
        return True
    if al == "scope & context" or tl == "scope & context":
        return True
    section_only = {
        "scope & context",
        "scope and context",
        "changed areas",
        "known dependencies",
        "acceptance criteria",
        "acceptance criterion",
        "test cases",
        "requirements",
        "validation criteria",
    }
    if al in section_only or tl in section_only:
        return True
    if al.startswith("scope &") and "context" in al:
        return True
    return False


def _is_changed_areas_source_line_junk(raw_line: str) -> bool:
    """Skip whole lines that are section titles or navigation, before dash/tab split."""
    s = _docx_compare_normalize(_normalize_line_for_parse(raw_line))
    if not s:
        return True
    if s in (
        "scope & context",
        "scope and context",
        "changed areas",
        "known dependencies",
        "acceptance criteria",
        "acceptance criterion",
    ):
        return True
    if re.fullmatch(r"scope\s*&\s*context\.?", s):
        return True
    if re.fullmatch(r"&\s*context\.?", s):
        return True
    return False


def _build_changed_areas(raw_lines: list[str]) -> list[dict]:
    out: list[dict] = []
    for line in raw_lines:
        s = _normalize_line_for_parse(line).strip()
        if not s or _is_changed_areas_source_line_junk(line):
            continue
        if "\t" in s:
            parts = s.split("\t", 1)
            out.append(
                {
                    "area": parts[0].strip(),
                    "type": parts[1].strip() if len(parts) > 1 else "",
                }
            )
            continue
        m = re.match(r"^(.+?)\s*[—–\-]\s*(.+)$", s)
        if m:
            out.append({"area": m.group(1).strip(), "type": m.group(2).strip()})
        else:
            out.append({"area": s, "type": ""})
    return [row for row in out if not _is_junk_changed_area_row(row)]


def _is_placeholder_or_empty_ac(row: dict) -> bool:
    """Do not treat DOCX fallback / template lines as real acceptance criteria."""
    raw = (row.get("text") or "").strip()
    text = raw.lower()
    if not text:
        return True
    collapsed = re.sub(r"\s+", " ", text).strip()
    if collapsed in (
        "n/a",
        "na",
        "tbd",
        "none",
        "—",
        "-",
        "–",
        "not applicable",
    ):
        return True
    snippets = (
        "no acceptance criteria provided",
        "no acceptance criteria",
        "(no acceptance criteria",
        "no ac provided",
        "acceptance criteria: none",
        "acceptance criteria: not provided",
        "acceptance criteria not defined",
        "no criteria defined",
    )
    if any(s in text for s in snippets):
        return True
    if re.match(
        r"^(no|none)\s+(acceptance\s+)?criteria(\s+defined|\s+listed|\s+specified)?\.?$",
        collapsed,
    ):
        return True
    if collapsed in ("context", "& context", "scope & context"):
        return True
    return False


def _build_acceptance_criteria(raw_lines: list[str]) -> list[dict]:
    rows: list[dict] = []
    ac_i = 0
    for line in raw_lines:
        s = _normalize_line_for_parse(line).strip()
        if not s:
            continue
        m = re.match(r"^(AC[\s-]*\d+)\s*[:\-–\.]?\s*(.+)$", s, re.IGNORECASE)
        if m:
            raw = re.sub(r"\s+", "", m.group(1).upper())
            num_m = re.search(r"(\d+)$", raw)
            aid = f"AC-{num_m.group(1)}" if num_m else f"AC-{len(rows) + 1}"
            rows.append({"id": aid, "text": m.group(2).strip(), "test_case_ids": []})
            continue
        m = re.match(r"^AC\s*(\d+)\b\s*[:\-–\.]?\s*(.*)$", s, re.IGNORECASE)
        if m:
            rows.append(
                {
                    "id": f"AC-{int(m.group(1))}",
                    "text": (m.group(2) or "").strip(),
                    "test_case_ids": [],
                }
            )
            continue
        m = re.match(
            r"^acceptance\s*(?:criterion|criteria)?\s*(\d+)\s*[:\-–\.]\s*(.+)$",
            s,
            re.IGNORECASE,
        )
        if m:
            rows.append(
                {
                    "id": f"AC-{int(m.group(1))}",
                    "text": m.group(2).strip(),
                    "test_case_ids": [],
                }
            )
            continue
        m = re.match(
            r"^criterion\s*(\d+)\s*[:\-–\.]\s*(.+)$",
            s,
            re.IGNORECASE,
        )
        if m:
            rows.append(
                {
                    "id": f"AC-{int(m.group(1))}",
                    "text": m.group(2).strip(),
                    "test_case_ids": [],
                }
            )
            continue
        m = re.match(
            r"^(REQ|FR|BR|NFR|SR|UC)\s*[\-\.]?\s*(\d+)\s*[:\-–\.]?\s*(.+)$",
            s,
            re.IGNORECASE,
        )
        if m:
            prefix = m.group(1).upper()
            num = m.group(2)
            rows.append(
                {
                    "id": f"{prefix}-{num}",
                    "text": m.group(3).strip(),
                    "test_case_ids": [],
                }
            )
            continue
        m = re.match(r"^(\d+)[\.\)]\s+(.+)$", s)
        if m:
            ac_i += 1
            rows.append(
                {
                    "id": f"AC-{ac_i}",
                    "text": m.group(2).strip(),
                    "test_case_ids": [],
                }
            )
            continue
        m = re.match(r"^\((\d+)\)\s+(.+)$", s)
        if m:
            ac_i = int(m.group(1))
            rows.append(
                {
                    "id": f"AC-{ac_i}",
                    "text": m.group(2).strip(),
                    "test_case_ids": [],
                }
            )
            continue
        if rows:
            rows[-1]["text"] = rows[-1]["text"] + " " + s
        else:
            ac_i += 1
            rows.append(
                {"id": f"AC-{ac_i}", "text": s, "test_case_ids": []}
            )
    return [r for r in rows if not _is_placeholder_or_empty_ac(r)]


def _link_ac_to_tc(acceptance_criteria: list, test_cases: list) -> None:
    """Best-effort: link AC-i to TC-i in order (fills empty test_case_ids only)."""
    if not acceptance_criteria or not test_cases:
        return
    n = min(len(acceptance_criteria), len(test_cases))
    if n == 0:
        return
    if len(acceptance_criteria) != len(test_cases):
        logger.debug(
            "docx _link_ac_to_tc: partial pairing first %d rows (AC=%d TC=%d)",
            n,
            len(acceptance_criteria),
            len(test_cases),
        )
    for i in range(n):
        ac = acceptance_criteria[i]
        if ac.get("test_case_ids"):
            continue
        tid = test_cases[i].get("id")
        if tid:
            ac["test_case_ids"] = [tid]
    if len(test_cases) == 1:
        tid = test_cases[0].get("id")
        if tid:
            for ac in acceptance_criteria:
                if ac.get("test_case_ids"):
                    continue
                ac["test_case_ids"] = [tid]


def parse_scenario_from_docx(uploaded_file: BinaryIO) -> dict:
    """
    Parse a .docx upload (Streamlit UploadedFile or any binary file-like with read()).

    Returns a scenario dict aligned with JSON scenarios. Raises ValueError on failure.
    """
    try:
        raw = uploaded_file.read()
    except Exception as e:
        raise ValueError(f"Could not read upload: {e}") from e

    if not raw:
        raise ValueError("Empty file.")

    if not raw.startswith(b"PK"):
        raise ValueError("Not a valid .docx (ZIP) file. Only .docx is supported.")

    name = getattr(uploaded_file, "name", "document.docx")
    ch = _content_hash(raw)
    slug = _slug_from_name(name, ch)
    out_dir = _IMPORT_ROOT / slug / "media"
    out_dir.mkdir(parents=True, exist_ok=True)

    try:
        doc = Document(io.BytesIO(raw))
    except Exception as e:
        raise ValueError(f"Could not open Word document: {e}") from e

    state = _DocxParseState()
    image_counter = [0]

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            paths = _save_paragraph_images(block, out_dir, image_counter)
            _register_figure_labels(state, paths, block.text or "")
            state.assign_images(paths)
            _process_paragraph(state, block)
        elif isinstance(block, Table):
            for row in block.rows:
                row_paths: list[str] = []
                for cell in row.cells:
                    for p in cell.paragraphs:
                        row_paths.extend(
                            _save_paragraph_images(p, out_dir, image_counter)
                        )
                cell_texts = _table_row_texts(row)
                merged = " | ".join(cell_texts)
                _register_figure_labels(state, row_paths, merged)
                state.assign_images(row_paths)
                _emit_table_row(state, cell_texts)

    _harvest_structure_from_notes_parts(state)

    # `_buf` is shared; only flush it into the section that currently owns it.
    if state.section == "story_description":
        state.flush_description_buffer("story_description")
    elif state.section == "business_goal":
        state.flush_description_buffer("business_goal")
    elif state.section == "acceptance":
        state.flush_description_buffer("acceptance")

    acceptance_criteria = _build_acceptance_criteria(state.acceptance_lines)
    acceptance_criteria = [
        r
        for r in acceptance_criteria
        if isinstance(r, dict) and not _is_placeholder_or_empty_ac(r)
    ]
    _link_ac_to_tc(acceptance_criteria, state.test_cases)
    acceptance_criteria = [
        r
        for r in acceptance_criteria
        if isinstance(r, dict) and not _is_placeholder_or_empty_ac(r)
    ]

    story_description = "\n\n".join(state.story_description_parts).strip()
    business_goal = "\n\n".join(state.business_goal_parts).strip()
    notes_extra = "\n".join(state.notes_parts).strip()
    wc_block = "\n".join(state.workflow_context_parts).strip()
    if wc_block:
        notes_extra = (
            f"Workflow context (extracted):\n{wc_block}"
            + (f"\n\n{notes_extra}" if notes_extra else "")
        ).strip()

    workflow_extra: list[str] = []
    _reconcile_figure_mappings_for_test_cases(
        state.test_cases,
        state.figure_registry,
        workflow_extra,
    )
    _append_registry_orphans_to_workflow(
        state.figure_registry,
        state.test_cases,
        workflow_extra,
    )

    workflow_process = list(
        dict.fromkeys(state.workflow_images + workflow_extra)
    )

    placed_media = _collect_placed_media_paths(state.test_cases, workflow_process)
    placed_fig_keys = _registry_figure_keys_for_placed_paths(
        state.figure_registry, placed_media
    )
    notes_extra = _filter_redundant_figure_caption_lines(notes_extra, placed_fig_keys)

    notes = (
        f"Imported from DOCX ({name}). "
        "Figure/Fig/[Screenshot N] captions register screenshots; steps may cite "
        "**see Figure 1.1** to map evidence. "
        "Otherwise images follow document order per step; overflow goes to workflow context."
    )
    if notes_extra:
        notes = notes + "\n\n" + notes_extra

    n_embedded = int(image_counter[0])
    ing_warn: list[str] = []
    if n_embedded == 0:
        ing_warn.append("No embedded images found in document body.")
    ing_warn.extend(
        _docx_structure_trust_warnings(
            acceptance_criteria,
            state.test_cases,
            business_goal,
            notes_extra,
        )
    )

    return {
        "scenario_id": f"docx_{slug}",
        "story_title": state.story_title or Path(name).stem,
        "story_description": story_description,
        "business_goal": business_goal or "",
        "workflow_name": Path(name).stem,
        "workflow_process_screenshots": workflow_process,
        "acceptance_criteria": acceptance_criteria,
        "test_cases": state.test_cases,
        "changed_areas": _build_changed_areas(state.changed_area_lines),
        "known_dependencies": _dedupe_preserve_order(
            [
                _normalize_line_for_parse(x).strip()
                for x in state.dependency_lines
                if _normalize_line_for_parse(x).strip()
            ]
        ),
        "notes": notes,
        "ingestion_meta": {
            "warnings": ing_warn,
            "images_detected": n_embedded,
        },
    }
