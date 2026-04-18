"""DOCX export: execution draft vs review export."""

from __future__ import annotations

import io

import pytest

pytest.importorskip("docx")
from docx import Document

from src.scenario_export_docx import (
    build_execution_draft_export_docx,
    build_uat_review_export_docx,
    safe_execution_draft_filename,
    safe_export_filename,
)


def _paragraph_text_joined(doc_bytes: bytes) -> str:
    doc = Document(io.BytesIO(doc_bytes))
    return "\n".join(p.text for p in doc.paragraphs)


def _paragraph_texts(doc_bytes: bytes) -> list[str]:
    doc = Document(io.BytesIO(doc_bytes))
    return [p.text for p in doc.paragraphs if p.text.strip()]


def test_execution_draft_docx_marked_and_has_placeholders() -> None:
    data = {
        "scenario_id": "SB_exec_draft_test",
        "scenario_title": "Login flow",
        "workflow_name": "Authentication",
        "business_goal": "Users can sign in.",
        "scenario_context": "Given an active account, the user opens the login page.",
        "acceptance_criteria": [{"id": "AC-01", "text": "Valid credentials grant access.", "test_case_ids": ["TC-01"]}],
        "test_cases": [
            {
                "id": "TC-01",
                "title": "Happy path login",
                "steps": ["Open login page.", "Enter valid credentials.", "Submit."],
                "notes": "Use test account from vault.",
            }
        ],
        "changed_areas": [{"area": "Auth UI", "type": "feature"}],
        "known_dependencies": ["IdP staging"],
        "notes": "Draft scenario note.",
    }
    raw = build_execution_draft_export_docx(data=data)
    assert raw[:2] == b"PK"
    text = _paragraph_text_joined(raw).lower()
    assert "execution draft" in text
    assert "pre-review" in text
    assert "pass / fail / blocked" in text
    assert "actual result" in text
    assert "evidence" in text
    assert "open login page" in text
    assert "draft scenario note" in text
    assert "per-step" not in text
    assert text.count("pass / fail / blocked") == 1
    lines = _paragraph_texts(raw)
    i_ctx = next(i for i, t in enumerate(lines) if t.lower() == "scenario context")
    i_changed = next(i for i, t in enumerate(lines) if t == "Changed Areas")
    i_deps = next(i for i, t in enumerate(lines) if t == "Known Dependencies")
    i_ac = next(i for i, t in enumerate(lines) if t.lower() == "acceptance criteria")
    assert i_ctx < i_changed < i_deps < i_ac
    assert "area: Auth UI" in text or "auth ui" in text
    assert "idp staging" in text


def test_execution_draft_strips_generic_required_field_steps_for_reply_draft_title() -> None:
    data = {
        "scenario_id": "SB_strip_gen",
        "scenario_title": "Reply draft",
        "scenario_context": "Context.",
        "acceptance_criteria": [{"id": "AC-01", "text": "Draft works.", "test_case_ids": ["TC-01"]}],
        "test_cases": [
            {
                "id": "TC-01",
                "title": "Reply Draft Editable Before Send",
                "steps": [
                    "Update all required fields with valid values.",
                    "Open the reply draft panel and edit body text before send.",
                ],
                "_export_primary_ac_slot": 0,
            }
        ],
    }
    raw = build_execution_draft_export_docx(data=data)
    text = _paragraph_text_joined(raw).lower()
    assert "update all required fields" not in text
    assert "reply draft panel" in text


def test_execution_draft_omits_changed_areas_and_deps_when_empty() -> None:
    data = {
        "scenario_id": "SB_empty_deps",
        "scenario_title": "T",
        "scenario_context": "Some context.",
        "acceptance_criteria": [],
        "test_cases": [],
        "changed_areas": [],
        "known_dependencies": ["  ", ""],
    }
    raw = build_execution_draft_export_docx(data=data)
    text = _paragraph_text_joined(raw).lower()
    assert "changed areas" not in text
    assert "known dependencies" not in text


def test_review_export_changed_areas_before_acceptance_criteria() -> None:
    data = {
        "scenario_id": "rev_order",
        "scenario_title": "T",
        "business_goal": "G",
        "workflow_name": "W",
        "changed_areas": ["Login modal"],
        "known_dependencies": ["SSO"],
        "acceptance_criteria": [{"id": "AC-01", "text": "OK", "test_case_ids": []}],
        "test_cases": [],
    }
    raw = build_uat_review_export_docx(
        data=data,
        test_cases=data["test_cases"],
        scenario_key="builder",
        traceability_matrix=[],
        session={},
        tc_to_explicit_acs={},
        gap_rows=[],
    )
    lines = _paragraph_texts(raw)
    i_changed = next(i for i, t in enumerate(lines) if t == "Changed Areas")
    i_deps = next(i for i, t in enumerate(lines) if t == "Known Dependencies")
    i_ac = next(i for i, t in enumerate(lines) if t == "Acceptance Criteria")
    assert i_changed < i_deps < i_ac


def test_execution_draft_filename_suffix() -> None:
    fn = safe_execution_draft_filename({"scenario_id": "My/Scenario!"})
    assert fn.endswith("_execution_draft.docx")
    assert "My" in fn and "Scenario" in fn


def test_review_export_filename_unchanged() -> None:
    assert safe_export_filename({"scenario_id": "X"}, "sk") == "uat_export_X.docx"


def test_review_export_still_builds() -> None:
    """Regression: final review export path unchanged."""
    data = {
        "scenario_id": "rev_test",
        "scenario_title": "T",
        "business_goal": "G",
        "workflow_name": "W",
        "acceptance_criteria": [],
        "test_cases": [{"id": "TC-01", "title": "One", "steps": ["a"]}],
    }
    raw = build_uat_review_export_docx(
        data=data,
        test_cases=data["test_cases"],
        scenario_key="builder",
        traceability_matrix=[],
        session={},
        tc_to_explicit_acs={},
        gap_rows=[],
    )
    assert raw[:2] == b"PK"
    assert "scope" in _paragraph_text_joined(raw).lower() or "acceptance" in _paragraph_text_joined(raw).lower()
