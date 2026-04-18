"""
Regression tests for JSON/DOCX import behavior and trust metadata.

Run from repo root:  pip install -r requirements-dev.txt && pytest
"""

from __future__ import annotations

import io
import json
import shutil
from pathlib import Path

import pytest

from src.docx_parser import parse_scenario_from_docx
from src.intake_parser import load_scenario
from src.scenario_media import resolve_media_path, step_texts

FIXTURES = Path(__file__).resolve().parent / "fixtures"
DOCX_DIR = Path(__file__).resolve().parent.parent / "test_assets" / "docx"


class TestJsonStructuredBaseline:
    def test_title_ac_count_steps_preserved(self, project_root: Path) -> None:
        path = project_root / "data" / "sample_login_flow.json"
        d = load_scenario(str(path))
        assert d.get("story_title") == "Improve invalid login handling"
        assert len(d.get("acceptance_criteria") or []) == 3
        tc01 = next(
            (t for t in (d.get("test_cases") or []) if t.get("id") == "TC-01"),
            None,
        )
        assert tc01 is not None
        steps = step_texts(tc01)
        assert len(steps) == 4
        assert "Open the login page" in steps[0]


class TestJsonLocalImagePath:
    def test_relative_path_resolves_and_label_kept(self, project_root: Path) -> None:
        path = project_root / "tests" / "fixtures" / "json_local_image" / "scenario.json"
        d = load_scenario(str(path))
        tc = (d.get("test_cases") or [])[0]
        ess = tc.get("expected_step_screenshots") or []
        assert len(ess) == 1
        assert isinstance(ess[0], dict)
        rel = ess[0].get("path", "")
        assert isinstance(rel, str)
        assert not rel.startswith("./")
        resolved = resolve_media_path(rel)
        assert resolved is not None
        assert resolved.is_file()
        assert ess[0].get("label") == "evidence"


class TestJsonMissingImage:
    def test_loads_with_warning_not_raise(self, project_root: Path) -> None:
        reg_dir = project_root / "tests" / "_regression_missing_image"
        reg_dir.mkdir(exist_ok=True)
        try:
            payload = {
                "scenario_id": "missing_shot",
                "story_title": "x",
                "acceptance_criteria": [],
                "test_cases": [
                    {
                        "id": "TC-01",
                        "text": "t",
                        "steps": ["a"],
                        "expected_step_screenshots": ["./no_such_file_ever.png"],
                    }
                ],
                "changed_areas": [],
                "known_dependencies": [],
            }
            fp = reg_dir / "scenario.json"
            fp.write_text(json.dumps(payload), encoding="utf-8")
            d = load_scenario(str(fp))
            assert d.get("scenario_id") == "missing_shot"
            meta = d.get("ingestion_meta")
            assert isinstance(meta, dict)
            warns = meta.get("warnings") or []
            assert isinstance(warns, list)
            joined = " ".join(str(w).lower() for w in warns)
            assert "not found" in joined or "referenced" in joined
        finally:
            shutil.rmtree(reg_dir, ignore_errors=True)


class TestDocxBankingEmbedded:
    @pytest.fixture
    def banking_bytes(self) -> bytes:
        p = DOCX_DIR / "doc_banking.docx"
        if not p.is_file():
            pytest.skip("test_assets/docx/doc_banking.docx not present")
        return p.read_bytes()

    def test_goal_clean_ac_tc_and_images(self, banking_bytes: bytes) -> None:
        d = parse_scenario_from_docx(io.BytesIO(banking_bytes))
        bg = (d.get("business_goal") or "").strip()
        assert "AC-" not in bg
        assert len(d.get("acceptance_criteria") or []) >= 2
        assert len(d.get("test_cases") or []) >= 1
        meta = d.get("ingestion_meta")
        assert isinstance(meta, dict)
        assert int(meta.get("images_detected") or 0) > 0


class TestDocxHealthcareSparse:
    @pytest.fixture
    def healthcare_bytes(self) -> bytes:
        p = DOCX_DIR / "doc_healthcare.docx"
        if not p.is_file():
            pytest.skip("test_assets/docx/doc_healthcare.docx not present")
        return p.read_bytes()

    def test_test_fallback_and_sparse_warning(self, healthcare_bytes: bytes) -> None:
        d = parse_scenario_from_docx(io.BytesIO(healthcare_bytes))
        tcs = d.get("test_cases") or []
        assert len(tcs) >= 1
        title = (tcs[0].get("text") or "").lower()
        assert "test" in title or "change" in title or "save" in title
        meta = d.get("ingestion_meta")
        assert isinstance(meta, dict)
        warns = [str(w).lower() for w in (meta.get("warnings") or [])]
        ac = d.get("acceptance_criteria") or []
        if len(ac) == 0:
            assert any("acceptance" in w for w in warns), (
                "trust signal expected when no acceptance criteria extracted"
            )


class TestDocxNoImages:
    def test_text_only_doc_images_zero_no_step_paths(self) -> None:
        pytest.importorskip("docx")
        from docx import Document

        doc = Document()
        doc.add_heading("Regression", level=1)
        doc.add_paragraph("Test: minimal text-only scenario")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        d = parse_scenario_from_docx(buf)
        meta = d.get("ingestion_meta")
        assert isinstance(meta, dict)
        assert int(meta.get("images_detected") or 0) == 0
        for tc in d.get("test_cases") or []:
            for item in tc.get("expected_step_screenshots") or []:
                if isinstance(item, str):
                    assert not item.strip()
                elif isinstance(item, dict):
                    assert not (item.get("path") or "").strip()
